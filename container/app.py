# app.py — Dual workflow: 1360×768 (work) + 1920×1080 (report)
import io, time, datetime, os, uuid, json
from pathlib import Path
import sys
from typing import List, Dict
from threading import Thread

# --- HELPERS: BACKGROUND ---
def run_in_bg(task, *args):
    Thread(target=task, args=args).start()

# Flask & Firebase
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, jsonify, make_response, session, send_file
try:
    import firebase_admin
    from firebase_admin import credentials, firestore, storage
except ImportError:
    firebase_admin = None
    firestore = None
    storage = None
    print("WARNING: Firebase Admin SDK not available (local mode)")
from werkzeug.utils import secure_filename

# Image & Docx
from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as Align
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Vertex AI & Google AI
try:
    import vertexai
    from vertexai.generative_models import GenerativeModel as VertexModel, Part as VertexPart
except ImportError:
    vertexai = None
    print("WARNING: Vertex AI not available (local mode)")

import google.generativeai as genai

APP_VERSION = os.environ.get("APP_VERSION", "2.66 (Final)")

# Logo PMA + Gen consolidat en BASE64 (Garantia total que mai faltarà)
LOGO_PMA_B64 = """iVBORw0KGgoAAAANSUhEUgAAAaEAAAB5CAYAAACQsp40AAAQAElEQVR4AeydCXxVxdXAzwtkIzv7bqjWrYBY1LqBC4oLLlhUpNAarICtglRF8BM/a8UFq34I1r2CVUQR6wYKAi5Q0CoomysoYQ8QCAkJWSHf/OdlHvfd3PuSvCxkue/3zp2ZM2dmzpx775xZzsyNKPV+ngQ8CXgS8CTgSeAISSBCvJ8nAU8CngQ8CXgSOEIS8JTQERK8V+wRlIBXtCcBTwL1RgKeEqo3t8JjxJOAJwFPAk1PAp4Sanr33KuxJwFPAk1PAvW2xp4Sqre3xmPMk4AnAU8CjV8CnhJq/PfYq6EnAU8CngTqrQQ8JVRvb03DZ8yrgScBTwKeBCqSgKeEKpKQF+9JwJOAJwFPArUmAU8J1ZpovYw9CXgSaHoS8GpcVQmErYRy57wj+6Y9K7gHd2UGlVuad0DsOAiKiorkUHaOFK5aKwc+/Ein3T9rjmS/NDMIyDNv/iIhDjryI70bEA+QtxXAAW7pwBNPGZSJa8qlbPjCJQ7IX7pcijdu0nUgrR0om/yseHCkRVbkbY3z/J4EPAl4EmjqEghbCWU/8Kg0a91KfPFxknnbXZL1yBNa8dBo7zhngOyZeF9AtjTce/82Wfb0/61sSU6SjJN7yu6L+smeawbK3t9dI/vShknB+4ukNCdXmsXFS+HqtZL3/nzxRUZKdK+e4otrEciLRp78UBCZd9wtOwcPl+1nXyRbOv5K5d1JQVIZtNY44jKu+p1ASxoUIIrBZEjelFFaXKyU4ltSsOxzXe7BzD2S//Y8zR98Arv6niXbf5Gq8u8kOy68UqgzvJi8MsfdLdQdpUMZxINDRr7EeNl3132G1HM9CXgS8CTgSUBJIGwlFNG6pewbf5+06HOmtHvleYlo11q2tmujG+0W118nEhOtGvV3ZNfI0brh3n/vJCn6dKUqMkWB/x87+nZpu2SZHFVaKi0f/qsUrV8vRd9+L4k3/F7aTn1U4q++Upq1bS0oHkYpO4eNkK0n/Ebnh+LKe+xpKZg9X0pWbdAZ+hJbiy8xtQw6aRxxhW9/LNDuSxuuFeC2E08XFJNRFpSRMORqafvcNF02PJRs2iytHn9I89ZmwWKBV52hvkRK0aIvJGf8RMXLL2XXmDt0XSMSEyX22oFauaJsI5RMyDP2jNMk556HpVlqF526hi9edp4EPAl4EmiwEghbCcVceL4c2rJOjzaYeku6fmigoU4YPEgSrvmtMHrIf36mEk57BblKOcQomhsCigdFE92zu1ZUWf83TVrf/7/S8n/HS2S3oxS9CKMWGvjNarTFyKlg5lxVZqbKxyiaJOWP0bSlOZlqJFWg/VxKcwpUOFN5CzWNLxHaTsqfKqW5+YJigj+UBcqtYMXXilZ02fCQcudfJOuhRzVvMb85RStFlCVKM3bEUJUP5eaqNK0kf9qLuq5xVw6Q+GsGKpxI3O3/I8gE2WxJPk7xvU5izj9Hx3kXTwKeBDwJeBLwSyDC71T9mpg2VCVqrxr6bMm46Cq9TpLyl9ES0fEEOZSXJ7svIz5e0WRJ9MDzpNUbb0vnzat1Yx6rRk8qQq/57Lzyd5L055EaH5GUCFpQCEyzMW1HA+9LNEonpqzx12T6grKJHXG1dN65UcFaXQbldN6p/AqXOPlezaPYfr5E8vLnWzDzTdl56q/1FFu+WveBlNERoxh4g0emGcHDO3jKoE7RF5+l0FkK4iXzuhuVK4rHVEm6eaSenkQ2Sp2KSHtBOSuP9/ck4EmguhLw0jcaCYSthGikW746TQkiV02HfSM01Hv+fLtEnXeWZJx9iWr409WoZ5S0/3qNtH/rVWFqzSiZ0rwDejqsZMs26fDJXLXu00PlI7rRZlSCQiiY/Y5uzEWiVV6HRzia0HYpPVAgRT/8KLkfLJCcf83SgL9w1Ro5tH+/jTo4iBKjDBRdkZpiY90n45KrtQEClNG9emge4RXFCO/gqQt1av/BHF3HuNvHyKHt3+m6R19+gSCLXdemadmI5Col/IyeWiStB54EPAl4EvAk4JdA2EqI5IUrmcIqVt4std4zT+L+cJ1eH0p54iFJmvqMHt1Eq0ZcEQT+rMNsPfYUSbn3Lkm589YAHusx1pQYlaAQUAylOenCSCNm6GWuiogRTdHHyyTnyeeleO23cihrnwb8+//5shx46XVVRrSC8n8U0OH80xVBoVZ8hfOXqbWeVNlzzyTBok9FaF6T77pN4N1u5UYdWz/6gCROniKtpj+pZRB73VVaJiJZIlIsBZ9/oVzv70nAk4AnAU8CVgmErYRY68h77EU11TZEWLhH6RSv/0k32iU7MmT/358ULOIwKACwFNtx7mXa8KDdxlWB0Q/WZdtPOU+yx9ykFYBV+bT78ithpBGl1o38U1pW1v3+0pxsiR1ylbR7fbqgCFjPAfCDS1bTcSJ7/MQO14hfpmql0fHndIkZOkgpO5SRaF5yJ/1Ndh7TW69NiSjOlELt/OMKKfxipVAX6kTdAOqa++QLakS23i+DDRu1IkY20QOHCLJCZuL9PAl4EvAk4EkgIIGwldD+19+UmGsv1lNtzdq20Upn/70TJPPMi5RC+at0XPGxtLi0v2Q/+LiwzhJzxqnS7p1XteFBVFSUZoARxfZfnC7FK7/XjT7GBREdU7RSQ/lEdu2iDQNyxj+s4pN0mqpeDhUVuiZhFJU/baYwzQYRVn4ovua9jtHKyJfYSQ5tyZSMk3sL5t3QYNKNkqMu1OnAh4sl5/+e0nWlztlj7tIyyJ10j5ZJ5C+P1jKKHnieIDPyqC/w3XffyWeffS7PPfec3HPPPTJhwgQNDz/8sLz++us6bvfu3fWFXY+PJiABnjcnyM3FCKgJCKAJVjFsJZQ8epQkjRuj13Z2nnOFaqy3SdQ5A7TyYfqM9ZCYU06WqD6nS/Lom4QFfdZRjIwZRWRecqEOogxQQKyrdNn2rTQ/92y9r4fpueadO0vc7TcopZCtaWvywnQcvMZceoGefsOcPEopjU5fL9VTa6U525Tyw4Chk+xLG6ZNsU351IU6YUUXpRQsdd156TUSO2KolkFk73O1THb0OldYY0q+e5ySwyiT/Ii56enpgpJp2bKlnHjiiXLmmWfIqFGjZNKkSTJ58mQNd911l1x33XU6rm3btnLKKafIjBkv6RHeEWPcK/iISCDcQlEaffueIwMGXBYE/ftfJO+9955jtkx987w5Ac+mYyIP2eAlELYSouY0vAfTtygFgeFAliTcMgK0xA8brNdDtnQ6UYqWfq5x1gsbR3PGj1UNPNZpMSp9tsRPHK+n0xgd7YiOloMZu6Tzzt165NRMNYSi1lWsedSMv1B8LZO0KXXX3Dxp1q6dbElO0lZ7rFclz5iueNumi2KdKn/aY4LhBC+LRpZdCj9aKpuTuqkR3ScSf/VVGps4Yaxys1T6AlWXnYKsFOKI/eld3nDDDdKtWzdByWRlsVYlEhHRLCTA8MqVK2X48DSJVveFUZO9/tB44EnALoGlS5fI++/PC4KFCz+UPXv22km9cBOWQER16o4xQcmq/yhlEqOyiZe4spFN8/btJOqCKyTxwbsldtDlEmE58YB9P3mPPajSpKo0h/+l+fmCIQCjo7ZLlul1mmZtWx8mqGUf02yt7p+oLd3YhIuy4QQHtRIUKBlFVDDzJcn6/ajAqMAXHaXrmDR1kh4JMv1GghYXnqecSFXPGEFGyEohjsifnie9y+nTp+vyrYpHI0Jc7LSMms4880zZts2vnEMk9aI8CZTr4FRWJPbnLlQ6ppWHDBkidLKsAI64UGnrMo6pb3iy8ogfXHp6el2yUq/KigiXG0yV8158RSVvr3r72WrKbEzgeB3MomOvuFSPMJi2o4FXhHp9h9EEjTlhA2wkRTEVr1glXdWIhGkuE1fXLpZuXTav1SMkv7FEsCJknahg9kzZOzhNs8a0HHVkYyoKl+N6iAAfO2KUlo1Ie0FWyIy4ugRGLldccYUu0rzYOhDGxaRnZNRZTZM25RcnDPF5SWpJAvv2Zctrr70mdLKsAO7gwYO1VGrVs83M3O3KZ77qhFc9x8aRImwlVLT+J93DZz1HJEvirhwQkEj+x0uFxjyAUB5GOfnPP6lGBsEjIBWlG+rYEbdoSzijsMAfSeA0Byz+WKuy84EiKnx7jp6as8ZRZyzlDC7h90OUN0vV2T8aKtlVt4v8rOMwclFM6B4prh0OHTooocBOTxhlhMvUHtN8+D3wJFCTErA+kxXlm5AQr0l4Lq2gkfXwYuURfz1ksU5ZClsJFaxeoxiNVAqkQCK6dJeYX/dSYf+fNRJrmNMGsBazj4D81FyLBQMEfPUJmndor9jJVVD+jyIqmPmCPsTUxEb37C4FCz8yQSHMCRIGcWDJf4y31t1169bpdRwKcnrQzUt+3HHHyUMPPSSLFi2Wb7/9VsPy5Z9p3FFHpQYUFPlYweQ5ZswYK9rzhycBL1UFEiguZj9iBURedIOUQNhK6OCmrarC9EAKJXrgRYGpOMyxI085KRAuWPG1PtTUl1h+BKQyaNB/6oSBBcYUVIQpuGYn/FKbpJtw9IB+SlFj2RcvfpkRU7uA4QDzzJRilAV+Aygg/O+++66sWbNGm2X363e+nHDCCRrOOON0jfvxxx9kypQpkGplpD2WC3kz5bF48WHFa4n2vJ4EqiQBtm5kZ2eLE/ztb3+rUl4eccORQNhKqGQrSoiKZknCNb/FoyFz6CjBbJkA6yO7+hHHiAJM4wMUEcYUbLqldq0mThBkgB8wU3IikXJYZsTUHsyYMUMYCaEk7KWggBjh8KJffvnlwotvpzFh4m699VY9SgJHWlw7jB17qx3lhT0JhCWBxMREcQKexbAy9BLVXwmUcRZR5lbDaS9m6g0LsIRxt4ixats98lY9CvCvG1WjiHqftL3sviZNc0ndkYHZ3OqXDSNGHV3rF0ZBf/3rfY7loERSUlLkyy+/0C+6I5EDklESIx6HKI1C4aWnp2t/ZS+5ublaUc6Y8ZLet8RmWYA9TG+++aZUNT/KJc+cnByxAzIh3gBheH7iiSeEMgHKxYqwJta4yB+rLDb8ki/5A/jBhVM3e51MmLJMvXDhH/lR1oQJEwQ/+FCA3OAXWqtMrDwTD12ofGoijjKoG64BwvZ6Egced/9+5ylz+CEOGmgBez7Q2AEa6svz4CQPnlmeH/Kzp7WHTdnkeeDAAXt0IIxxhaElX+gDkY3cU00llKVPTcCYgFEPFmBYiiEzzogrmD1DLcon+RVRUt01xJRfF+CLj/XXLTFG7xFCCVMuMsid8oxgDYdsogdertDuL4qKrLH/ggULZMeO7eUMEVBAFPLGG3OkTZs2eKsEgwcPFtaPyMcKJhMs5ow/lEsjSePWvXsP6dGjh163Yt8Sm2UB/FdffbXez2Q2yYbKzxp36aUD5PjjT5CePU8KQFJSUtDmSBranj176rLHjh2rN+macrEixJQds9lwTNBN3Y499ji9EZgNv9SH/AH84DDoAhQncwAAEABJREFUqGrdLrvs8grrhiUk/CM/ymKD57x586wiCvIbfrkXbFwmnVUmVp6Jhw7FFo5sggp2CdDwdu3aVd87yjKQmpoqjz/+eFAq6gYemkGDBgXFWQPEQdNTPRP2Z8FKhx8lgPLuqZ4P6svz4CQP9szx7FI+8aE6FYMHX6fvG8/En/70J4pxBDufK1d+5UjXGJHVVEIicdcM1HLJvOdv0nr6U9qPQtr7x7HK31430rEjhkqL6wdrv0I2ij+nLUSd/RtJuG+8qlemUrapkj3mjsDp262enyp7H/a/OHHXMiWpFlZjomu97q+++qprGTTSjGpcCSqIeOCBB6R3797Sp0/fIAC3NTA9654JowAaSRq3TZvStaJkytAJyAXFxgufmtpN6JmCCwVZWXu1AiZvA9CXlJTgyNixY4WG9ocffnAtG0LMfDFBr8paF8otnLo1a9a8RuqGcrBaQiJT6uIG1K2y/Jq8kCmNP7LhXrrlXR18VlaWUI4VwOXl5QVli6ECeOicOl0Qwzdx0ADgzLOA3wrs4UlISNAbuXk+iCN9KKB8Rkp0KnBJY4f9+/cHnknoyc9OA87Op52mMYfDV0JlDWps37P8DW9BYcAse+/9k8sa5hglu2hp/fcHlNv4/qVFhfpEh4guqaq+nBoRKXsn/FVXlBMSDu7cqT9PEXP6qRpXF5dQ02a33XZbtVigt7ZixQpZsuTTIADH2lGozOlhMgqAhpcOwO8GxBugAaFnSkPhRm/Fm3S44Fu0aCGUT0MBDgDvBMQBxF1wQT89ZYg/FJA3yg0a0gL43YB4A4wqq1s3FArKweTpVq7BQ0/dCFc2jaHDJR33coaaSsVfG0A5BkLlXxGNicd1y4epNY6vIh46A4RDgZWOERHPgRu9ldaJxsTjOsU3Zlz4SkhJxZeYqtd/cl58WVLuukNhRCskNp76Elurhjldkmc8KliNFcxdoOJrfySgCqmTP+tcxV+s0mW1euEJ5Wao0VCSMAWJhaBCCB/Ey37mn/prrSIpoGoVTCNtf5Bp6Cj41FNPwalzoNfM9BAFO/EGf1aAzgomDQ0FU0jWuMr4reVDby0LPzg7mDLT0tLsUUFhRkA1VbdQ0zpBhVoCO3bskD/+8Y8WTGgv6w39+/fXRKaOOqAuyMIJVFTQ36RjlBrO/QjKrJoBw69bNiYe142mb9++OsrUSwfKLqSzQ1lUwDHpeA7sI/YDB/wjOJNHIJHNY+JxbVHVDtb3DKqlhGKuH6TXPUo2bipraKVsJJCiFFCBNO91tj41gUX6ksDxPvVdJJXn79D2DfpzFS36ny/RFytZ5DAaai97x4zXmUT36iHF33yv/UxJak8tXrZv3+aae/fu3atkjOCaURUjaKToNZPMvKz4AesLx5QePBq8NQ6cSZuWNpxgpYF0M2fO1NNvJLLna3BOeNIyJWiUO7RWoG7WEZA1zppfZes2cuQoaxaV8mOEwkgRXk0CygZM2OoyUibOSk88ONyjjkpV63TDhZGtuW/EAcQbMOnfeustg6pTt0OHDsL0MjwyPWznD2bAEQcNtBde2F9atw5eD8X4wGmajLRAhw4dg+RBvcEDlGEAPP6XXnoJJwBnn312gE+eA3s6CMHBm5XPhLINuMQ3dghfCanpt9gzTpecOW9J8h1jtJwKV63VIwGO4RHJkJZTJwvrQ/vSbhLOkmMdRRM2kktk77Nk/72T9eiv1VOPqVoxGorRJ0mYkxOSxo0RNutGn3W6iq/d/6pVq1wL6Nevn2tcbUawjkT+5iXFD/DioXTYGFtaWipM6a1du1YKCwvl2WefhcRxbxIHYropBZ3I4ULZlEcUee/atUsoE8A/ZcoUohzLI+LTTz/BKQePPcY99x8Ca42kLBqvqtQNHjnck6kya16h/KTZUWaEQpkGQqV58MGHykWTDiT8pqdvlBdffFGQyaxZs7ScpkyZQrSjfJYsWaLj6vqCkpw3b67A41NP/cO1+GeffUbTQPvhhwukX7/zg2inTXsyKEzAyIN9dHTsrPI4eLBEH71jpcNv4F//etl4tYvsKBs+J09+ROOcLk88MSWIT94NJ7rGiAtfCSlpRJ1wnOS/PS+wFrTv0akKm6JHQXzKILbPmZKlcHyiIVr53T5MpxI1uD8KtdnRqZI09SHZe88kPRKMGXqjqjsbU9tL9t+f0HVibShv9r8l+lcn6nBtXraFOFSUdYfKlM10DVMK6enpkl5JMPSB/Ms85PXEE345lKG0Y17yZcuWCRtjNbLswn6QkSNHCg0AKEOLn0YX9/nnn8OpNJg8UHLkbbUOxE+DRmPhluFHH31cLoq6sQ5jj6AszOBXr17lWjdOp7CnM2G7FZjBu7nIhDKJpx4oEoxEUK6AlUdGbps2pQdGhaQxwAjJfi9MHPkykjBhq5uVtc8arHf+gyHOjuMerljxpSPPPA/so3OKxFJ06NChTlHaCMExQiHN1JzylvuH4rMccSNDhK2EmrVrJ0Xf/SBx12L55V8LKpj5ml4XYRSU8n8P6FEQx/WwebW0ER67gWFCi1HDheN72Kza8v6J6vHIUjKIkaJF70qhGhkqhJZRSfpmaZaUTLDWYOfOXa55Y57qGmmJWLhwoTYvxuKnsoCCG+kwlfTpp5/qnGkotcdymTNnTsjpQRoAtxcd6zUaEEt2FXqfVaOrUL1LYz5rGnRrhps3b7IGtT9U3Z5++umQZvD0xocPH+44smCkh7LQhVTiYvhF+dBwokg6deqky0fBAiabNWvWai9prADyqqv8nyDB7wQ9enR3QjdoHJ02puKohF0eV155JWhXuOCCC13jvIiqSSB8JdS2jRxQo6D4Ky7VJe5/+TXtMkKwjoJiR9+ujRf4VIMmaEyXgiJ94kDi5Cl6LSyy21ESc22aGg2xNpQi2S/654eRUd4bbys5BM9H16UomjdvXqniGIlUitBGFBkZacOIOK0X8LJDiJLBDQU333xLuWij0L755ttycU4IU15FjSz1Zj3EKY8CNfVsxzOKs+MoC/7OPz94ysdOR3jEiJE4QUBaEBs2/IRTaUDBonwqStCnz9nC6MgJqL9bevbPLFq0yC26weKPPvpoV3mkpqaGrBdTdDYCLximBMJWQpQX1buXPiOOTZn5b81VKDakZkjc6BuVX4RRUNIN12t/Y74kpg3Va2GsfyX88feqqqwNJUn+tJl6NMiG1chfHa/wtftPTExwLcBtf4RbAhrEyoJbHm6NKYuwpGFzYihITk6CzHHEsHYtB+jq6AovbLK1jgjcEnTt2tUtqhx+7dp15XAgzjrrLD0KCVUv4rp27QJ5teqG0iOTtLQ0nAoBRYMcnMCamJEYU6ysT2F2zMZaDDR4Hqx0Dd1fFXmkq6lp5MH0MvJYunSJ47RmQ5fJkeA/bCV0KGufxJzh3/9S8NUqvRiP2bJIe0kYeLnkznlHGBFF9+pxJOpVp2VyVA/rQTn/miVYykV06V42GsqQAwv96wmx5/URZFabjLVTU6Ru+WdkZLhF1Qqe3vPmzZsd82YeHkshdqaHAjN6cWr89u7d65i3E7Jt23ZO6LBxKJH58+c7pl+3bq0+jSFUvYjrV2Yo4lQ3TjN3zNwBifUXjalDVKVQKBvMzMeOHas/w00D27ZtWz0ly14izI7ZvOnEZ6UKaGBEKBss5iZMmBCQB50YpqaRB3JqjAr5SN6msJUQTEcddyyO5L0zT7kpuuGNGXqZ3heU/cCj4v/EtYpqAv+ksX+WnIkP65rG3ThMuYUK2kvuzNeVK2JkpQM1dLFnw4tix5nw6tWrjTekSwMLAb1sNyC+IuAjXTk5GGkEU9KYMQ/PBkEat4oA+uAc/KGqKNWYmBh/ohq8Iht7dvBaE3XLzMy0Z+0a/vWvT3aNCxVBr75jx05a2WBmTg+f9SgaWJOO+gCEneoLvrEAFpcoYN4hjuvBoMPIg3tKPZGFgcYuD+pbVxC2Eor85dFqjcP/1dG8x55Wi/FMnWRIwh+G6FMC2BcUdcUldVWPI14OVnB8AK9w1VqJ0+tkWUomMVL49iy9l4rREjKrTUZ/8YtfuGbvZOHlREyvb+LEifp7QnxnyA5MpVX2BczOznEqQk9jmJe5ItcxA4WsqmGCSlIn/4rqY413Y6g2LaXoZHA2Hr16zLvhwcqT1c99BgyN8RNuTMBZhmyENgrYKgOrn/obwAISf2OSw5GqS9hKyBxFw/eCRLIC/Dc/92zJ/WCBsEBfnWmCQIYNyJNw30TJe/f9gMm6YT1/2efaa2SmA7Vw6dHDfeoT01ysgSoqFguy+++/X5iOcILOnTtVlIWOj46OlqSkRHH68fJWB8gzlCUg8UcKqlMv0sJ3Ts5+nFqBO++8U7AutDauFETZdgCPeTaWd6NHlzcSIb7uoWZLZM2Lswzt8qAUJ3kMHz5cf9qEdNB4UH0JhK2EsASj+ILPvlBOvJ6KY0MqiufAK7MDB5uqyCbzb3Fpf8l71m8RF3f7/yiZMB2VIrlz39cyMDLTgVq4pKam6lx5ebSn7MILhnfhwupZOOXk5OgGjLwqAqbAMBW208EbLzIN26JFi/ULHY7LaM2ed12FecadLOmoG7ve62vdmHJi2s08D0Ze8I2fDsiUKVP0PeF7U2zmJVwZyzvSNzRgPYw1LzvfyANgwzH15/nEohB5YBXXr9/5EhXVeI4gs9e/rsNhKyHDaP6b7ylvvII90mLYtXrqiT0yLS48T+Ga1j/6hOOEo3wO7sqU2AuoPyPEaDUlx7l5dSOL8ePHuxb08MMPCdMxrgQVRDz1lP+UdHsj5pQsKipKOnTo6BSl9wfRsPEyhwukd8y8jpChyicu3HqRjvS1UQ16/G75sll17dq1+rgeeODDcm60jQVvjtixPs8oH+rHNDSnJTASRB5YFIL3oOYlUC0lhGl20acr9dqHSLHEnNRTCr/7QZ8Zx6GlNc/uEcmx0oViih1z7VDJX7JMYs84TafDYvDQlm36aB+NqOXLsGEYRZQvhBcNI4Cq7sg3OTGV59RrNPFO7oknnuCElrfffscR35CQxxxzjCO7X3/9dbUUvWOmNYCk88FCuz0rGl1GppwCYI9rzGHkschl7xOnQzAVHar+2dn1+6SIULzXt7hqKaGi9Wyq2xOoE+bYTM9F9/OfShuIaEKe6LNPl4LPv9AWgodNtbOkZKv74aI1KR6mVDDbpXFxyhdFMqOKR/Dn5ubKxRdfrLNDmWlPJS79ysyQraSkZ32KqSEr3snP9B90ToBSdEpTV7jTTvN3Muzloegr80Gyuq7b9u3b7awGwhdddFHA7+ZZvHixW1SDxBcUFIgxRLBXgE29dpw9/PLLwWfE2eO9cOUlUC0lVJKxU5VUrNY+CiTqnAHKL1Lw/kKJOd35BdUEjfwS2+dMOfC8/8Ny0Refq2qLqXZ84Agfhaj1P+alFGJXRCgA8MOHpwk9PRpCwqEAU+ru3Xvo7+qY9KHorXGnn+4/tNXOBzQjRy38NgAAABAASURBVI7ACQljx44VrJacoDK8h8y8mpFOdTPyYRGfnnaoIkLVzZgEh0pf1ThM5quUxkLMpzB4Dkz9LFFN0kunyE2BNUmBVLPS1VJCRWu+UcWzHlQokaecpPwihfPflOjevbS/KV6YhizNSddVj9RTNsXKHylF69crt27+rCkwl+1UGg0JgKJKSWmplREvVXp6urBTnsadUQb7SDDlxeKOkQtpnPILhWMe3en8N/KiURsw4DKhXHseNOCYzWLFRRz0Bgiz+H/CCc5TfcTXBVA3JxnDJw0UZ4u51Q3LKuoGrRXgm1Fs9+41f05bbGws2TvC22+/7YgHyXOAvPE3RFixYmWV2X7//Q9c0/Cu0CmCgHuHWxOwcePGmsimQeZRTSXE0SUooWJpflTXwFx4RFxcgxRGTTDdvC3nw7XXa0CRv0hVWeYqiJZDW9ynQxRBjf8feeQRCbWXgReIEQrKiJeKTXrsEUpNTRU+38w+EhpKGIMWFyANgL8yYBbD7WnIkzUKyu3f/yKtDGmcUXypqd3EpIPOlGPyGDPmVoM6oi6jGRgwfOEH4JljXdzqxpQoNNDagfthx9VEOFXdV6d84AOjBJQ+nRBomH6lseVe8ByAgw63oQGjfjo71GXIkCGCUqUOGF5ceGF/vEFAPbl30JtOBJ0i5MH95l0hAXS4NQVskLXySSetpvKu7/lUSwmV7sH6iyrmChsxfdt2SETHEwKbWImpPDQOSowTInsfr9eArCPCg+lbtOVgXdUS6zTWJypSRLxMAHwxDQTgB2eAMEBjO3z4cMFsFT+4ioDGD3ro7GnIH/zChR8KjS+NM4qPTZTEAcRbAaslRnpW3JHyUzcOD3Uq3/DuVjd7GmSD2Xlt1s1p5AYf8IrS57gen88nCQkJehqUe0E8AH+4dti/v/b2NNnLcgvzjBNn55F6gaezQ11QtpmZu0FpGDDAf/iyDlgupIOeToTP5xP2vKF8MG83ZPayDD6Ua/+gnqGlPPxWPvfvp/MKtvFDRLhVxDLuoF4T8ucQddwvdcMb0baVH9GErxFdOsihvANiRoRYyB38eavG1aVYmDLC7JYpHl4awK18XgQr2OlIy7QM+yROO+039uiQYRo/ACLywTVgLdPqN/G4pAFQgKxlgasvMHLkSMEsHv4AK1/W+lj9Vhr8pKNubBImXFswbtw4nTXlaY/lYuXP6ocEU3uUv1M6ThNnChe6IwU855xF6FS+vS5WGmSOAnOqlzWd1W/SO8kDOuKZzsa1Q+/ev9aoisrTRE3oEr4SKimRQ7v2lJlni25waXibpfpPB25CMixX1eZHd5Pi9T+JLzpKySdVx5fm1I11nC7McmHD6JIlnwq9bNC8AAD+igA6AzS0fB2SNMmW062JB1cRMBoCoCMNgD8UQANAw0uPAsRfEZDGQEW0TvEmLa5TvB3HNKIZEZHGgJ3OGjY0uHVVN54F68cCKdvKk9VPHEAjvXz5MjHf1wFnAHr8jBrwOwHxVnCiccJVJQ2j/smTH9HZVCUdU3JLliwJSqcDDheTL1F07G677Ta8+hR0E4cLctq0aTjlAD6nT5+h8dAa0IgmfAlfCRUWyaHtZjouRYuwZOdOiWjjjYRYHzu4a7f4mjcXXxJrZoinWI2E8vAcEaCXzeKndURiXgI3F0YZRS1f/pnQ0BIGwjUKoOyq8EBZGDaQZsKECQTrLTAi4oum1NEw6SZX8NBQN07Mrsu68R0nGtHevXvDQrlGFN4AIlGOnFbOtGOXLl30GiN4O4wdO1ZYM7Hj6zLcr9/5wnPK82rKpR5WAJ+fz7e+8PkBIxCeL7M+ZKW3+qGmI8b0I2lQKLjg7cDUMidx2/GE09Ku1ydSGPmDs5aDH1xTgrCVkF9I/nlLX2KS7vWX5uRKRKLzeWF++qZx9SXGy8HsffpbSxGtkrUJe32oOY0Jo5Hs7Gz9IvBSMcXGC4FRAi8VLzGNI3Q0kIyinNYpOMbECrNm+c3SK6qn4YGXmUaDEQR8MDUC0IjT+C1atFjg85VXXhHSVJQv8R999JGgCKx8EZ458xWiK4S5c99z/MjZmjWVO4GckQZyq0zd4JG6VVahUzfS2IG1nAorZiPgPq9YsUJQRvCL3LnnuEb21MGqHOPj44U1RuRp5wFc77KpJlMU9HY6E/7d74YYsiCXht3Q2F2ekSBihwDPKc8raakbz5cdnNaBeL4+/HCB8LzzPCIHIw9mEBg9kicdMeplimYqkroTZwVwofZeoTCRP2nc+PzVr040xTRCN7hKYSuhQ3l5KqdCBYf/paqX4SkhkWZx8VK6N/uwYMp8fPSuzHtEHaYheBF4qWbNmqV6kMuFlwLTYl5iGkeUQagGknl4K5BnVSrFy0yjwQgCPphqA2gUafzgr6p5wg+KANeACVeGN8oz6awu+MqkNzSVqRv5G/rKuNA7AWVVJr0TDcqI+4zcuee4RvZO+VK+kSd+A+BQIPYyTLzddaI1ae20JuzEj0ljd0lD3Xi+7ECcnd6Eed55HpGDkQczCIwendLxXFB34qwALlQdTXmkceOTvA1dY3fDVkJ+wbAHRsQXH6unng7t3y++Wvh2izSwny8+TqS4pMFwzQvDS47bYJj2GPUk4EmgUUigmkooUguhNDdfSktKJCIhQVBEGtmEL3xB9fBaUBMWRP2puseJJwFPAvVUAmErIb/5cbSuVmmOf7HPFxsjpfn5GteUL4eKCvXaGJvcDu3ZFxAFpykEAp7Hk4AnAU8CngQkbCWE7HyJ/i+riuRKaWGRsCB/KCeHqCYNh3Zmii82ViKZkjt4UMnF/3lpv+Ju0qLxKu9JwJNAXUmggZQTthLSe2CS1AK8HgXlavPj5u3aycGNmxtI1WuPzeING6T5Md2UTA6UM2OvvVK9nD0JeBLwJNDwJBC+EmrePOh4Hiy/mrVuJYcy9zY8KdQwx4e27JDIDpyagAVhhs49omN7bcauA97Fk4AnAU8CngS0BMJXQnEthONpdC7qUrzhZ2nWqqWUrNpQp2ekqaLr3Z/jjCJapgiKGeZYM2vWoZ22ICTceMGrmScBTwKeBKomgbCVEMX4T0dgr1C8FP+cLpHdjlLoDCnZtVu5TfOP4ilZ9R+J7NpFClavUUKIVCDS7OhUvXlVB7yLJwFPAp4EPAloCVRLCUX14Lsn7BWKFNZBdI7SPjAC8Ieb1vXgXv9RRs3atpaDm7aqynNsT6E069JR+b2/JwFPAo1NAl59qieBaimh6F49VOkc3RMtxV+uUn6RmKGXlY0AdLDJXUq2bpPmvc7W9S7+5nvlYsaeJZE9ms4xHKrS3t+TgCcBTwKVkkC1lFDzzp1UIfHaBJkpqNK8AxJz4blSuOxzhW+a/wMffiSxV12mK18we76WDYGYk3rieFCPJfDdd98JB0/yOevnnntOAPzgiKvHrHuseRJosBKolhJiDcgX2CskUvDVKmnR92wpeH1u4CurDVYyYTJe9PmXEnteHyneuEnl4LeMEzVF2bxjBxWuoX8tZsNpyJyKzBliVgC3e3fjW+sz9fX5fHLiiSfKFVdcIRzqOmrUKAHwgyPO5/PJDTfcIIsXf9Rkn++aevTS09OFZ8r6jOEH5yn8mpJyw8inWkqIKkZffG7ZKdHxUvDZl9o4oTQnXXzbdhDdpACjhKJF70rMr3tJgVJGIpFaNpG9jw8yZ6/PQtm8eZPwBUmOo7cCuDx9aG195r7yvNHQDRhwmVi/mMlHyUIBufPtnAsu6Ce9e/fWygicB1WXQG5urutztm9f+cN/q16Cl6KhSCCiuoy2GDhAZUGPP15yn3xB+VkXulEOLPmP9jelS/5nX0hk73O1FVze7H+rqrdSkCux1w5UbsP62xvjhsV9aG5RqIxs+JyytZ6hU4lYadetWycoo7Fjx1aUrM7j33zzTRk2bJgetTFyswJHSdU5QyEKtMoUfwhSL6qRSiBsJZS/dLkWSczpp2rXlxgjh7ask4O7MiV+2GDJm+ZXSDqyiVzy/v2OtLj+Ol3bwrdnla0HKSV0fl+Nyy+TmQ54lyMiAZQGQOHVafRMWhQaI6r61LjzPZuZM2cKozY7UG8PPAnUJwmErYSKft6o68G6UESX7nraSdT0U+4HC6RFnzOleOUnWiFpoiZwoRHKf/5JibvsYrErm5hTTtZrCIWr1jYBSdTfKvLdIpQGCgRw4pQvWzqBEy15AIyobrrpJieSI4qDNwNHlBGvcE8CISQQESIuZFTJho0BJRN/y42KFlPtVnLgldl6Oip64BDZ//qbCt9w/1XhvPijJdo0G6Wc8+x0lbS9UszZEjv6duUXvUbGJx7E+x0RCWBMcNddd7mWbRQPBBgjPPTQQzJlyhTho28pKSli4nGhsQINPSOOGTNesqI9vycBTwKVkEDYSsgXGSn5S5bpIuKvYc0jV08/sTCPZVjy3eMkZ+LDOr4pXLInPiBJd9+hN+oWzJyrZSGSJQnX/FZXH0MF74N/WhR1fmERnPUbCkZh4FoBxQL+3XffldLSUpk1a5ZgqYUCQhHt3btXNm7cKMOHD9fJoNce22X48DTJ8U6Rt0nFC3oSCC2BsJVQREqyFHz+hc49sttREnXBFarnX6DC8bL/5deEKajSnEyxT00pgkb3L1TTbEw/xl99pex/+z1VvwwtC6YpsZRTCMn/+FOJaNcarwd1LIEnn3xSl4ii0R7LBYXSp09fycjYIXzG2RIV5E1NTRU++4xSCoooC5i8n3rqqTJMQ3U8vj0J1K0EwlZCpfkFwsfbMESA5RbDrlVOhhoBtJbcx55RfpHkGc9I9oOPaz/f19GeRnjJfuo5SbjPP+rLnULd26ta5grTlL64FnrPUERiolJMuQp/ZP6YJLPpkjUR1kaYOmKPTG1yk56ers2YKctaLlNjxNVm2SZv1urcpuFQQEy1zZ37nrRp08YkCekyOrruuuv09JwT4YwZM/T6n1OcE477gjzYFGtkhEsY/LZt25yShcTFxsa6xkdFRbnGhYrAGhB+ZqgpR/jjGcKFT56jxriHLJQ8nOK4V8gIuSAfNjsTBu9EX99w3Ef4hW9TB1zaDeKYUagNnsNWQsVbt0j8ZZdKzoyZmq+EgZcrBZSq/ewTyn5ppiRdP1QK578rhWqkoCMa2yUmSq+LYZCQcvto4bQETo7wJcbomsbraUrRI8MW/ftJyabNGl+XFx6g1NRugY2YY8eOFRplpo7YI+Pz+fR+jZriiUafh7hjx07SrVs3bcZMWdZymRojDr5oxGqqbKd8FixYoNFmpKIDlsu8ee9LouogWFAVelkvggglhmuFH374QX766ScrqpyfKTsaKZ/Pp+8L8kCxGRnhEgbfuXNnadmypb5HyLZcZmUIn88nPp8f2GRbhg44pv4+n5/G5/O76aqjECCyeVAs99xzj863R48e5e6l4ZPnqG3btsL9REmF4tNWRJWCNIQ+n59vny/YRUlWlBkNrM8XnM7n84fd0tNspNrEAAAQAElEQVQI+3x+Gp/vsNu37zmB4pAhFpLcK+4ZcuEd4z4QBt+//0XiVgZy9vkO5+3zBfvpqAQKC+HBNN/nC07r8/nD1MMpKXljwu/z+fSeOfiFb1MHXDZrc48TEhLklFNO0aeKOOUVLi5sJSQFhcLnCorWrNNl8+nq+NtvUr39TK2M9qXdofHJM6ZL1l33ab/LpcGifVHRsmfSw3oUxIgna/y9qi5lBgkj0vTGXV7I4h83iC/Wr5gUQZ386bXwYvAAbdqUrsukIbIDETxovFSkad68OaiwgEaiffv2+qSBHTu26zzs5ZkwkfBFY3v88ce7vqDQVQfclBwKhA2nZ5xxepWzT01NlQsv7K/TkY8BjVCX//73C3V1/tMQJiUl6Y6AoTAycXKhycrKEu5RdHS0IGNwFQF52WmccHYaE0ZuKJZJkyYZVNA+KfKyAkTcTzocPANuDS50jQm4n3SosJCkXlaZGD/4hQs/FBS5kzJgFM6UMHROsHTpUid0Ody8efPK4Qyif3//82rCuHSE2C+HUQ1hw6+bC83KlSv1qSK0FzU1wosg43Ch4LMvJHHUcMmd847OIvnWPynXTDlliBkNFX/zgxTMXaCUU5KKbxx/RjtFn62QgpfeFEZBefMXyeFRUJak3OVXwoVvvitJY/+sT5Ooq5qj+M4991zhxbA+UE7lm/ilS5fI4MHXSUlJiRNZhTh6YfSWaDBNnrhuCYkzwOiBF5QX2o0+HDxK9T//8RvPOKUfN26cE7pSuMcff0zmzJlTDl577TXVW+ztmAf1o6dJpKk7LmE3IN4ANMg4lCKy0kJvh4rioUcB0TnAb+hxCbsB8QZ4BriftaWITDnGdePJDW/SGdeNzoo3tLjgGRUwirDfT+LsQBoAPJ0JRov4rXDzzX/WQeisAPKdd97FqRCclAmJ6GydcMIJeAOA4Q2jNRCmPPyhwNDh0l507XpUjRjiVE8JffSpxPY5U3L/+bLmndEQayOlOduUwkmVfWo0xKGmrWc+qxtoTdSILoe2rJOUZ/6uTdL3jvyLqpl/FBQz9EY9ClIIyfvXa9pIo/DjyvVmSFNduP7664UeCw+LW1723ju0KK0//YmOhFsqZzyN4tVXX60jyUd7LBdTlnEtUdpr0vBC82JrZA1cMjMzhd65PSv4AHf++efjhAXdu3eXQYMGlYPBgwcLcfZMmYKjfuB1ffGUAfw4QVl0wDHpUEQBpM1j8rGhddDE4WqEwyU9PV2fnUeUKQ8/QDonIM4KJl1aWlqV1sesedR3P8db/f73v9dsmvpaZaMjbBdDx2iRKThr9IABnDwjjmuNvJd2emta/EbhmzLAGbjlltHGq106QxzJRcBOb62D1Q+tFUhH/J//7Fee1riq+qulhArfnqXLi73uKr0eQoDRkC+xE14FubLn3ge0oooeOERN1RUoXOP4l+YU6H1BCUOuln3TntWnRTA6EsmSlvdP1JXMVSPEuD9cp/2F8+tmzxRrQPTGeUh0wbYLDw5g0PgNkCZLTf3gmviKXEZdZ599tiazp7Pmy3QDAKHB4zdg0poX2+Cr4/700886uclbB8ouxx13nDAtVhasdefvf/+7LsPOC7IgAtkw+gAuvXSAnvoyccQbMOm5zwaHy/Qg6UiPErSnNWForBAZGUnyABjrP1OOiTDp7XwSb+LwGyA9HaGVK78yqEbjUjcafeqHn/oDVLBDh444WpkYnEaUXaDH+89//hMnAPHx8YEtAAGk8hj6NWtCb3RfvHixonb+DxhwaVCEddRljTD88iwNHTpUd0a439CYOPxW4GSObWEYz1jzqJYSIqPCVWsl8eqrJOf//KapjIaSpk5SCiddjYZaS95jD2rDhKiTT1LkfIVVOY3k3yy1izZMyB5zh6prqqrzNr0+hMk6Vdz/7IuCkipY8TXBOoFRo25yLYcHiQeM/TAc7QKsXbtWJk70K03izUPvmokt4umnn9YvnD0deUHKlFVW1l5ZsuRTDbt27ZLp02cQpdNpT9mFPHix6amVoarl/PTTBtf0LLCGaynmmmmICOvaiiEzMlq+/DMtG/YnAfPmzRVkhiWeoTFpjLts2TLj1e6HHy4Q0pF+9Ojgnq8mKLtAY4VOnUyHUYTe9rRpfnP2MnLtwAONK88K95EyAPLhfg4fPrzcvdQJ1eWLL/6rro3vz7MKIBsabGSDLL7//jvhvaIjSK2Jx7XDI488YkfJ0KHDyuEM4r33Qk/JvfyyfzbK0ONSNu87a06EAZQn09/wTtiAoWU/HM/SK6+8Itxj7jc4p46NyWPz5i0mm4pcx/hqK6GC1Wv0dBSjIdaAKAWrOA7yZLQg0l723DhGSjL5DAAfeIOikUBBkeweNUZVJl4poAKJ6HKcXh9SCD06Shh1A161HvSFdjFp155aujAthkGAeTisxfCQoWx4wNgPQ+MD8HDdf//9wksEPXS4lQFGQU8rJWSnNXmQJ1NWVuszXoi0tOuFhteezoTvu+8+462Wu3fvXtf0WCy5RtZwBC++yRLZGABHh8DJOAKZMSrhXkIPrRU2bdpkDQb58/Pzg8LWAPfMGrb6N2z4KaBMKNMANKz58azgtwL385lnnnEduYXixZpPQ/QjHywlabCRDbLgvvFeMS3L80+9oMO1QpaacbDfi1NPPUXYMuBEP3v2G65Tm3Qe6LxZ8zf+P/7R3waZ8IoVK7WXMqwAcs6cNyQ1NRVvEICbPXt2EM4a2L59mzVYZX/YSqj0AFNr8YGDSlE87JHhcwZw0eYNersZaoQQI8Urv5f8aTO1n7jGAEy9Fc5fJoVvf1xWrwxh7QsrOfZO5b34irB5lbriF4nHW6vA3LFTATxsDKtRNk7x4HiJGLXgryww1eLUqyI9Lyd54ncCGt7x48cHGj0rDYue1R3ik1+oT0/Yp6Ggry3YuNF/zqJT/hdddJETOoAbMsQ/nRtAlHlycvaX+WrOcRu1HHVUqnC/3EpiRHnyyb0co0PdA8cEDQTJO8W0Jgv8bizz/PPeucWjPKxxKLA//OEPVpT20xGhc/njjz/qsP1ipuqgM3Hwh/+SSy7BCcCiRQsDfquHaVzKt+Ks/qOPPloHTb46UHYJ15ipLLmErYT8GUQqBbNCT0kRbvX8VNl98+149cJ84uQpaoTAtFxMWUOtoxrNBUUEcDIEZ8RhpEHldqvRUevp/ulJFBJWcyLBc+/Q1TR88MEHrlmahUhXAhXBqMWtJ6aiy/2dGi3zkF555ZXl6O0IDCjsOPMiff/9D/aosML1IdE555wj3377rT76B4VkBRpwNx7pKb///vtu0TWO/93vfufI4yeffByyLBpTt554yIQNPPKmm0ZVWINBg/zHdlVIWEaAhSpe8x7hN+C27uM2Vcc0qV2xMC1sff6Mf+rUqaYYR9coQPN+OhKFiYwIM534Whpz61wxZ8hxVI8Ul+g1IPJNufPWsuN8Gu9HqphybN6rlyQ/+iBVlvyly9X0ZJxE9+qhwwcWf6JdLhFRtTcdiTnyzz/7F+Ipy4B5mEP1ZA0t7s0334xTKVi9erUjHYv+dpNQJ0JoeKgNj1aaL7/0T2FacVX1x8XFuSYpLi52javpCBoC6pqamqqnO1Itrr2s9PR0vReIDb89e/ZUa0NZeqrLTlcbYaaTrLxZ/fby4JO1O/a9nHrqaTqae6k9jfxinlfM0CuqalJSckUkQfGh3tMXXnghiJYAHRU35XTVVVdBEgTWe2r1c++DCFUAS1XuMfuJKlNXlSSsf9hKqHnrNqpAXuT2kjvzdeX3/1s+/FfJuu8hf0Bd2815Wa2VdFIjIqbvFKLR/Qul7b9fEdOjzX7wcWn9+OH6Y6Ital1MpFiatW1ba7XHHJl5ZqcCGGo74Z1wp53mb1Cc4uy4NWvW2FE6zBRds2bN9U57n8+/Y9vnc3ZJ4NR4mZ4X8eFCqONrCgqO7PPIdON7772nT0K44YYbhFMRfD6fPmUCE2x2rSNHJ9mEK49w0sEn60EomyFDhkizsvvKBk1Mztn3smlTep0pynDqUFtpQnVyqlMmU9n29DwHrC0y6rTGbd++vdxG70OHDuq1JUbgVlo3P8qGvWEoGza4+3z+d5WNrNxj634itzyqgw9bCUWdeLwqN1dPs7EuwrSTQuhpOE4SMGGs5dq++6qKylDQuP4cT9RmwRxdZ2rG6eG+VilK2bQmqKcpC+d/pGUkkitRPX+l8bVxCbUAfNJJJ1W6SBqXyhI7jbx4WUhveov4w4GdO3eFkywoDTv3gxCWwPr17pZzFrIa99Kz5EXHMILTLGjE2WRoOhDIzwo1zkAlM0RB9u9/kcAne8DgE4sv7quVP+OvZLYeWSUkEGoq+/PPPw/K4cMPP9Rh7oP2lF0GDhwY8jgqRlBsmuW0EpQNHVWUjVlXJj87lGVd407YSii6dy8LMxlipuRAxlx6gRR8tQqvhmg1NdVmwWI1GkrX4fKXeCn67EthY2v5uCOHOfDhIlV4ioLyfxRQ8oxXpEX/8wORBZ9/KS0GHF5o9k/FZQXizYnaAUQ99LBfobJsZWfnOJLaH95QYccMagjJjm63rHjZeBHd4msDz4iHniVlk7+TXGjk7QBtXQEyYcSDguSoGcq18wnOziNh8B5UXwJM3TKl7SRTRizWEj7+2Hm9LpS5N6PbM888U4YPTxMz2j6S9zhsJcRemMOfbwiekos5qacUrw8+xJHGutUbbzsqIhb3i5avlM3xyYH1JKug69rPiGZL1x5S8PpcNYpJKlc8Cihp6jOCRaA1smTLNok64bgAykzFsW4UPVBNZbT1j5ACBA3c065dO1ODIJeXpzpAZjt3Vn/k3LVrF7IqZ4HHC0dEdab8MIenQXACpk3I3wooIEY84CgfwA9YZUUY4xAsrxYtWqwaCvc9ONDWNFx11W+FEQ/8GTBlOPFp9shwNAzxhtZzqycBt5NLMFRh/Zfc6TBwr/AbMPegX7/DnWMTh0taOmcYkoS6vyYf9oexX40tFdAbPHnVFISthGAgadytyslQDXWMFL49S3+yQCH0P3vMXbLjwitl+ynn6WkpkJgst3z1jXKKqDQnWxLuHittl3wiGSf3lL1/m3zERkXsddr+i1RJvH+CtPznFMVrsA28UUDJo/2WMUw7UkfqmjPev+mTuqLIOCUBBSuSIUm33QK61oDFb7fMMzIy3KLK4XlIyyFdEMccc0y5GB5SNsjRgLIHJlyYPLn8Zr5yhVWAYL8GjaMbGQrELa4i/KRJk/SOcqYx7GCfSmQKzqqATN7ICsCMlz1cyAwrOu4XG0FpSEJNs5p8aspleoZRGo2NNU94BOCT9QrDJ/uwzB6ZwsJCaxLPX00JmCk55G7NimnbtWvXadTSpf/Rrv1+8SzpCIcLceRpTwMOcp5l9qehdLCcYw8Q4R49upfrzEFfE1AtJcTopnmvs1VDjfVbvOTNna95KljNgnWWRJ7UXZqf2VvjzIUTBFp/sFClSVdgkj51AAAAEABJREFUFoejhdO4MXHuvHO3FH29Wo2KOguHgpKuNk+gNhZrWLVtO7mP3vfU8ed0PcopXM1RGfGwoAEFlKym4IwC0kh1oY7UVdS6j7/uyqdGfSLxqo7ZEnXOAH10kSKttT8NLr0WpwJmz37DCe2I48FzjHBAHnOMf++AQ5TQgLIpNlwgvVO+VcVdcMEFrklQJPQmXQlcIpjOoLEmmpfZAGGgZ0+/ZSR+4PHHH8cJWrw3Lz09WXals4eLOjMVY4xcdKI6vDz22KPlSjN80pmAT/bFGD7LEddzxL59WTXLYS3mlpqaGjip3RTDc4b/00/9FrefuJjOs1EWOjtg1ICBiR3PPSZvNtdySgIjH6z04MFOWxvhiHAzZQSQecfdEnfDMIm++HyVTaHsf/gJPRrKnfKMJKnpqtaPPiDNkpJVXPA/7uILpN2XX2kkoyBGCwUz3xSmwEq275D2b72q4hfL3j+OERRDwUefKtrDykAFaugfLSXf/CAZl1wtu/peJMkP3SsdV3wsEXFxehSXO2myGuUlKUVSoCBdUJ72KbhDeXkS2bmLUFfqnDftBS2DnPEPKx4LJXrgxRI76HJBVshMIWvt7zQy4eFioxsNZ2UKfuuttypDpmmOPfZY7dovTEdVZURlT1+T4SuvHKiz40XTnrILcsFL449bFXA7B448mEazmrvy4huFRbwV2Kzr1mAYutrYlGrytro8H9w3IxdrHD1hOhNWnN1flybv9rIJ20ef4Oywfv16O6peh9m35cTg++/79wP+4x//CIrmGWctqXv37kF4E1hTdv6c0z3GwMEtnUlfW27YSmj/62/KwS3bJfEPQ6TlO69K7IhRUpqbL1kP+XtT0b16yK4xd+hPOOy6Nk1ohPnom6kIe4o671yrFRgjDF9ikhzakqmn4/bcM0mifnm0dNn2rVYMB9OrdzaRKdPuovyKlRLikNGjSvdL5Pl9JeuRJ2RruzZStOgLpYBaK+WzTZr3OkYYHaE8yYPeM6M06pc5/GY5MGuOrh91Lj1YEpABMkGhIiNklfvBApLXGoTaGDdt2rQKy922bZuYaaMKiRXBOeecq67Bfx5wlF6ob+qYFMiR/SZOQONt6Krj0qPjxXTLg9EQlmBu8XY8tE888UTQqMZKc9ttt1mDYk4MQC5BESpgFKTyuv7dFJhrgjAjvg+xObhfv34hc+VescAdkqgGIhMS3Dui69atrbCEqVMrfgcqzKQOCczBoygXa7GcKMIUL1NzVjz+tLQ0HEdwOkvR5P2b34TempGenq7zdHqOdUQ1LmEroaQbfi98T2dLcpLsu+V2zQJrKG2fmyYpk++TrL/cLSl/GS0dPpkn7ZSSSrp5pGA9tl2tERWs+FrTN1ML9e0/mKMP/UQRgfQlpgojkC3Jx2mF0KLPmXq0pSa4iHaEUGeylebkqjTOm0QZhbW4frA+ZJS1oJ3H9Jac8WOV8klVaUQpoHSJu32ctPnv4oAZNtN2u3/TTwq/WKnrR92oY8Kw63SdW06dLMggecajIjHRsmvkaEFGJT9u0FN8OuOwL6EThmosODGBnm6oHO644w4dXdkH7Ve/OlHvRzAPsk5cdhk//k7Xs67KSORxNU2FSbgTmF6boa2O+49/POWY3NQTSzA2hzoSWZAzZrwk0FpQAS8ywKCgXz9mBQJoqc6aDuWRk+ETf23BgQN5YWf96qtswRBXxRx2xraEzZo10xhkrT2Wy9NPPx3yeWMPDJ2jupClha1qeRlRX3qp/xMPJiPDP+8XOBM2MjFrScTZgTU8O66y4QceeKCypFWmC1sJcUZawjgW2+Ml//k5Cp6UA0uWagZoqPU6SbejhH1CANZ0Lf93vLR7/w3J/vtUwfhAE6sL+PZfr5HI3sfrhp+1FEYpLPRv6fgr4ew1X6KzZRkjKM6l2+RrKZuTuukpPab18G/ydRAMJEKlPfDS6zrNvrRheiSGEkQhRnRMkbZLlulpNjNHz2gu59npWrHCM3WibgCjIOqcX/bdoMJln0v+tMeUXOaoGkZK/NiblFu7/+5qGM66kHkgTWnmQT3ppF56N77BG5ee7JAhQwJWUQZfkYsxxMiRI8uRUR7WN5dddrmQdzkChcAogH0JyivQGyDs1JiDDxdQDBhL2OVCfpSLy+ZQTtaGL5Q1fAP4wfXo0UMwaYXWpMFvhTfe4F5bMSKhNsyGGuXQ0zXlBefoDyUmJvg9VbxSJ6ckLVrEOaE1zm1HPpHIhj1E+GsbWPfk2bCXw/1gJIZln71+hOHPPGv2tPU97HQ0EPXl/cK18o8RDmuKVpzVz4Zoaxi/ycMYO4Czwz333FPltsGeR6hw2EqITA9l7VMOZ6JlSPzE+9UU2i9lk88nxeu+1Q3wJl+UHs1Y10IY/bR7fbpEqum2HedeptdPVCb6mBvWY5LVwn9EF6bBGP75h98Hf94KiSMwmom7/Qbpsi9dOqz6RNgYC+Dvsu8HSZ4xTSm2TMe0IEuzcwVA+YjsAaXXs5gKxFACROGqtdrKL+b006TdK88LSgc8gBUcU3jUG6WDUQX+qB7d1SjqfxQJlmnRcminOw+KqMb+zz77jGNePGw0wuzGT03tJjxYAJ/p5TPOLJBD45g4BNJ8NoC8rWTkxT4T8h4w4DJ9MgCjDXqklI8VDvTQ4VrhpZdesgZrxD99+os6HzufIOEB4MWGLxQOfAP4waGMoAFIYwXyZDEXZWfF42dXPWmgIWyFSZMmCTJhWhJgHY1yMOdmPxG0pMW1Q04VDzA1+XTtepRgXMCUIvefRpq87cYU4AzQiKNs4M/wiYk6HRdkA53JH39Ng8mPTo/bga+Uj1LnnvF88czhEqauJg+n+2Di6qNr6mvnm/ra+XXqEFppevToaQ0G+WkXuKfcXyAnJ0eYeqZjxnPqVJ7JID/fGJgZTNXcsJUQo539905QpWXpxrbV/RO1Iono0l011C+otZ5BKi5eTW9N1GssO4eNENZRSvMOKLzoKbC2s2fI3gl/1YrK4Fn47/zdf5XyeEWMMmKfjU7keImW4tXr5MDCj6Vk6zYpLSnRgB9cwUIsSfzKzCk5eTPyIS5JzRl3Wv+lGOs3eGLEtu+hx/UIDhNz6DgpnLrsHDxcMOdmxCaSIuwFavfKC3o6L+aM0/QoKnb07SpJrpLDWDHTkApRa38WkBnC2x9aCuRBAjZtShceLID5ZXAAaQBoKwv0Tp999llHcvIkgsaBhozRBj1SyicOIN4AZdOoUQeDqykXPjE7JT/KwbUD/IQCOz1h8mKU5fR9GOKZUrn22mvwBgHlgEAm0dHRwjlxCQkJgtKbPn26Hh0ST/64dti/3/0UbbfzyiiT/Jia5X5w/82aFXxigk28tSzSEOa+wJ/hk0bL2nGxpyMNEIpP4qsKN944wjUJvAI8XzxzuIQBEjGKAtx4haa+AbMww4cPD8mWqU///v1D0vVQHWMIDD1+wMiHe8opI8cee5zwwUemnumYEU8aAHo7ZGfvs6OqFA5bCeW9M08XFDviFt3YEmAd6NCWdVK0/idJvv9uhcpSwJpMeymYOVcyL7lQth57il4nyV+6XB9vw6go5oxTJWPAtcK6jEogTPWhjNptWCltFixWCu0sNZpJL4MC5R7WvEzbFS36QvZcM1B29T1Ldp76aw34wVEuNORrwK94lMLKSZeoc3oLe5e6bP9GKx9GatDBCzzFntdH4BE8PLPGs+3E03VdCmbPV6TtFVDHLEm59y4p/O4HxV+6ZN1xj8KLtJ36qMQMvVH7jcx0oBYvM2ZM17m7PTQ8VFaA2NCyD8T4wVcG6IExEiAdYE1jLcfqt9KQBmA64fnnn7dG1agfIwWrIqLMcAsgLUDD/e9/vxk4O9Apv7vv5l2QcvssjDxIw3QSrsGRN2EjV/xW2LBhgzUY5DeLzCYPa6TJ3+Csn7S49957NdqezprGiU8+9eDGZ6ipPF1YFS+MNpG5nUdrNoZfXPCGtjZG2ORf2zA0xMfuTNm8O6mpqSbo6HIaCu+3UySyAjB2QHlDQxhAfnRs6YjgJ84KVdnWYU1n/GErocLFS/QIiEV4RgaMCiLatFL5Rkr2iy9J1l1/la65eVqJsNYjwrRUvBzanqHWSWZohcG0FY16qRrOtZ83WyI7dBBGTPtnzdEbXOkFsBcJ44Uu+7KFExdihl5WNkLK1I09n1EQYaMcysAJChWdod0mvvhYNWI5T420pgt7kjp8MlePylB8TBuifHYOGyFRv+gm8ARv8Kh5VUou//kZug6iP82QoZUYa0fwt3fcRMl57kUlg3hhbxN1QTZM4TEiKloafO6TIqyVP73aXbt2BYwGnB4ca8Emnu8JnXOO3+INnBWs9E7+KVOmyBQFxJl0+EOBlY7e3vLly4UXJVSa6sahiJANLxV5WXkgXBFY6dn4x96Zinhmrc6MFklvL4MX3QDxAA07fA4efJ2dXIdZZGe6RAdsF9YFUAqgyQu3MkADbxopp3SGR1ziARq/NWtWSz8HCzromF5kmsetfPIw4EZjxzPKQaaVSQcN6dlge9555+lTyQmDN0C4tsGUhVvVsvr0OTvwLrulHTdunFtUEJ6pWKNMnHjhnhkgHoCezdNOW0DI/F//+pdgWYs/HIgIJ5FJc0jNG9JA8w2h5LtuE5RF553bJSIqWloMu1aPaFAirPVgeJA4eZIaFQySyN6nKEWCLXuKUkhPyu6L+snm+DjhBOqont2l6NvvZff1NwVGTMxRsg7DdBgNOtN1HX9epw0HktRII+72P0nsiKsl5tqL9agp+uKztKIBRxzlMqLq+PN66fzjCsFsmpEWoxvyxnQchUGZxWu/FXjY97fJmid4y3/+SVXlFInoeILmPWboILVuNE2oE0osts+Zep0o9opLpVm7dkq5bdSySL5jjCAbZCQFShkqpawyqpM/imjz5s1CQ0mBPExuQEOyfPlnMmjQIDlwwD9dSpqqAg0fvSIUCmndyjN4aFAGNBAvvvhiyNEEtDUFyIaXinKZSiNfw1NFLrTUj5MN7r//foKVAkaLbPhEuYQqg8xQBD/++IPAp9kQbE0DDcB0idvLT4cAMIYq1vT4SQ/Y9/fQSDHNZqatoHUC0pL/ihUr9EGZPXr0AKVHe1Z6kEzzmLUnwm5AOrc4Kx6l/9lnn8n48eM1mnRuwP1lE2a/fudLOJaKJo01f11oJS9FRXSQK0nsQkZnnOfHKRq+wNs/XgfODWbNmiU8Y8ST3g3Y2sCzAD205oR9Kz14Rk8cdIs/HIgIJxFpOnwyT1pNulcDigHrMPA07Hrj5vVDCQaA+JQ7b1XrRc9Lh0/nCYqky7501WDvDgBrRIlpQyV59E3S9rUXdd5Rx5XfEMmoBcs0Gv/k0aP0dCAjMqbNUIQa3npVm0rDC+WiDElD2gBTZZ7oXj315xcoM+XOvwg8tJ35zwBfjJi6wKtSYPBOfSmXOpVloR1wrdTaGDIAQTy0Rk7tF78Dus6Al4JOphoAAARASURBVJWGMjs7WxjloJDo1fBi0pDyIKJ8aEgYIcBYH9XrogfuBBUN90kPDQqFMmnkaahQTqZczhqDj+nTZ8jWrVsFZUADQdq6Bsrlc+coTkYqyAQlYeeDXjcNHi8kcqF+jDbsdBWFWetCuSAXZIBMmFrCJcw9Qm4oAhoe8kMRUaYbsNYFnRMgd8pDzpSJEqQOlIMfXMeOHU3SgMsGWjow0MAXHQWeGfjkmQEPn+RvEnHf3XgETz0MLe6xxx4r4J2gd+9fQxISeLYxcqFuPEvcH/gEuI88dygf7i/3j8zgwak8cPADjR1uueUWVz7Jz05vD7PhlPydINS9s+czbNgwjUIBaI/lwn3BaMOCqtDLM8Y9tMqOe4zsuMe0C2vWrBGeBZMZz69TPQzO0FXVDVsJMTKhsQWqWiiKALDmQT5WsMaZF7Kq5VSGnrxNudYyDc64xMEzUJl8rTTWPKz4uvLzgDLKQSHRq+HFpCHlQTTKx/CCPHi5nMDQVMalTBp5GioaBFMuZ43BR1ra9VKVl7AyZYZLQwNKTxOZpKdvlNLS0iCgMaPB44VELuGWQzrki1yQATJZsuRTwSXMPUJu0FmBMt3ASufkp7FGzpRJI0IdKAc/OPhxSwcNfNFR4JmBT54Z8NXlk3Ld6kScE09OOOrGs8T9gU+A+8hzZ5SPNV1Vy0R+bmms+br5q5ve5MvI1PjtLp/asOMqE+YeWmXHPUZ23GPaBaf74CYL8JUp04kmbCXklJmH8yTgScCTgCeBmpfAV1/5jzmLiGgWyNyMivr27RvANUTPEVVCDVFgHs+eBDwJeBKoawlMm8a6dPlSmT6rziikfI51j/GUUN3L3CvRk4AnAU8CrhLAiCM3N1fYMIp1IZuC2fjtNAoaMWKkaz4NJcJTQg3lTnl8NhIJeNXwJBBaAhhZsDGYDaNYHWJMYlVAJjXWa6zdmHBDdT0l1FDvnMe3JwFPAo1SAhgMUDEUjwHCdjCfFLHjG1rYU0IN7Y55/HoS8CTQpCWAQQKWf1gpNhRBhOLTU0KhpOPFeRLwJOBJoB5IAMVjAHPt+fM5MqweMFYDLHhKqAaE6GXhScCTgCeBmpIABgnkZZQOLmE2UrN5mPP72B8FrjGAp4Qaw1306lBeAh7Gk0ADlcDUqVOF0ww4iYATIfADnH7B5uGGbpJtvy2eErJLxAt7EvAk4EngCEqAUxYwTkDZMOLBDzidYHAE2ayxoj0lVGOi9DLyJOBJwJPAEZVAgyzcU0IN8rZ5THsS8CTgSaBxSMBTQo3jPnq18CTgScCTQIOUgKeEGuRtqz9Me5x4EvAk4EmgOhLwlFB1pOel9STgScCTgCeBaknAU0LVEp+X2JOAJ4GmJwGvxjUpAU8J1aQ0vbw8CXgS8CTgSaBKEvCUUJXE5RF7EvAk4EnAk0BNSsBTQjUpzdrLy8vZk4AnAU8CjVIC/w8AAP//xIXNsgAAAAZJREFUAwCKKp1OVNdmxwAAAABJRU5ErkJggg=="""
def resource_path(relative_path):
    try:
        base_path = Path(sys._MEIPASS)  # type: ignore[attr-defined]
    except Exception:
        base_path = Path(__file__).resolve().parent
    return base_path / relative_path

# --- CONFIGURACIÓ BASICA ---
IS_PROD = os.environ.get("K_SERVICE") is not None
BASE_DIR    = resource_path(".")
if IS_PROD:
    UPLOAD_DIR = Path("/tmp/uploads")
else:
    UPLOAD_DIR = BASE_DIR / "uploads"

MASTER_DIR  = UPLOAD_DIR / "master"
WORK_DIR    = UPLOAD_DIR / "work"
EDITED_DIR  = UPLOAD_DIR / "edited"
SESSIONS_DIR = UPLOAD_DIR / "sessions"
REPORTS_DIR = UPLOAD_DIR / "reports"
STATIC_DIR  = BASE_DIR / "static"

for d in (UPLOAD_DIR, MASTER_DIR, WORK_DIR, EDITED_DIR, SESSIONS_DIR, REPORTS_DIR):
    d.mkdir(parents=True, exist_ok=True)

# --- FIREBASE INIT ---
BUCKET_NAME = f"{os.environ.get('PROJECT_ID', 'infofoto-vector-art')}.firebasestorage.app"

if firebase_admin and not firebase_admin._apps:
    try:
        firebase_admin.initialize_app(options={'storageBucket': BUCKET_NAME})
    except Exception as e:
        print(f"ADVERTÈNCIA: Firebase init error: {e}")

db = firestore.client(database_id='infofotovector') if (firestore and firebase_admin and firebase_admin._apps) else None

def get_bucket():
    if not (storage and firebase_admin and firebase_admin._apps): return None
    try: return storage.bucket(name=BUCKET_NAME)
    except: return None

# --- HELPERS: STORAGE ---
def storage_save(local_path: Path, storage_rel_path: str):
    if not IS_PROD: return
    bucket = get_bucket()
    if bucket:
        try:
            blob = bucket.blob(storage_rel_path)
            blob.upload_from_filename(str(local_path))
        except Exception as e: print(f"Error upload {storage_rel_path}: {e}")

def storage_download(storage_rel_path: str, local_path: Path):
    if not IS_PROD: return False
    bucket = get_bucket()
    if bucket:
        try:
            blob = bucket.blob(storage_rel_path)
            if blob.exists():
                blob.download_to_filename(str(local_path))
                return True
        except: pass
    return False

# --- FLASK APP INIT ---
app = Flask(__name__, template_folder=str(BASE_DIR/"templates"), static_folder=str(BASE_DIR/"static"))
app.secret_key = os.environ.get("SECRET_KEY", "clau-secreta-pro-2025")
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
app.config['SESSION_COOKIE_NAME'] = '__session'

@app.context_processor
def inject_version():
    return dict(app_version=APP_VERSION)

# --- GCS SESSION MANAGER ---
def get_sid():
    # Prioritzem SID de la URL per persistència robusta
    sid = request.args.get('sid') or session.get('sid')
    if not sid:
        sid = str(uuid.uuid4())
    session['sid'] = sid
    return sid

def load_gcs_session(sid=None):
    if not sid: sid = get_sid()
    local_path = SESSIONS_DIR / f"{sid}.json"
    if IS_PROD: storage_download(f"sessions/{sid}.json", local_path)
    if local_path.exists():
        try:
            with open(local_path, 'r') as f: return json.load(f)
        except: pass
    return {}

def save_gcs_session(data, sid=None):
    if not sid: sid = get_sid()
    local_path = SESSIONS_DIR / f"{sid}.json"
    with open(local_path, 'w') as f: json.dump(data, f)
    storage_save(local_path, f"sessions/{sid}.json")

def update_gcs_session(updates, sid=None):
    data = load_gcs_session(sid)
    data.update(updates)
    save_gcs_session(data, sid)

def set_vertical_alignment(section, align='center'):
    sectPr = section._sectPr
    vAlign = sectPr.xpath('./w:vAlign')
    if not vAlign:
        vAlign = OxmlElement('w:vAlign')
        sectPr.append(vAlign)
    else:
        vAlign = vAlign[0]
    vAlign.set(qn('w:val'), align)

def add_field(p, c):
    # MS Word requereix una estructura molt específica per als camps (Field)
    # begin -> instrText -> separate -> [text] -> end
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = c; run._r.append(i)
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate'); run._r.append(f2)
    t = OxmlElement('w:t'); t.text = "1"; run._r.append(t)
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end'); run._r.append(f3)

from docx.enum.text import WD_TAB_ALIGNMENT

def create_footer(sec, nat_code: str, dil_code: str = "", qualitat: str = "atenea"):
    # Netejar footer existents
    ft = sec.footer
    ft.is_linked_to_previous = False
    
    # Esborrar paràgrafs anteriors de manera segura
    for p in list(ft.paragraphs):
        if p._element.getparent():
            p._element.getparent().remove(p._element)
    
    # Crear nou paràgraf
    p = ft.add_paragraph()
    # Format de tabuladors: Un a la dreta per la paginació
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    
    # Construir text: NAT XXXX/XX - ART MNORD - Dil. XXXX
    parts_line = []
    if nat_code: parts_line.append(f"NAT {nat_code}")
    parts_line.append("ART MNORD")
    if dil_code: parts_line.append(f"Dil. {dil_code}")
    
    full_text = " - ".join(parts_line)
    
    # Afegeix contingut
    r1 = p.add_run(full_text)
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Arial"
    
    p.add_run("\t") # Salt a la dreta per alinear paginació
    
    # Paginació (mantenim mida una mica més petita o igual, usuari ha demanat especificament NAT i Dil a 12)
    # Assumim tot el footer a 12 per coherència o 10 per paginació?
    # "peu de document amb la font Arial 12 per el nat i dilig" -> Paginació no especificada, la deixem a 10 o la pugem a 12?
    # Millor tot a 12 per coherència visual.
    p_pag = p.add_run("Pàg. ")
    p_pag.font.size = Pt(12); p_pag.font.name = "Arial"
    add_field(p, 'PAGE')
    p.add_run(" / ").font.size = Pt(12); p.runs[-1].font.name = "Arial"
    add_field(p, 'NUMPAGES')

def add_logo_to_header(header):
    try:
        header.is_linked_to_previous = False
        # Netejar contingut previ
        for p in list(header.paragraphs):
            if p._element.getparent():
                p._element.getparent().remove(p._element)
        
        # Taula per controlar la posició exacta sense afectar el marge del document
        # Fem servir una taula 1x1
        tbl = header.add_table(rows=1, cols=1, width=Inches(6.5))
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = False
        
        cell = tbl.rows[0].cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Eliminar marges de cel·la per pujar al màxim
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), '0')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

        p = cell.paragraphs[0]
        p.alignment = Align.LEFT
        # Reduir espai abans/després del paràgraf
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        run = p.add_run()
        
        # PRIORITAT ABSOLUTA: logo_definitive.jpg (o png)
        logo_path_jpg = STATIC_DIR / "logo_definitive.jpg"
        logo_path_png = STATIC_DIR / "logo_definitive.png"
        
        target_logo = None
        if logo_path_jpg.exists(): target_logo = str(logo_path_jpg)
        elif logo_path_png.exists(): target_logo = str(logo_path_png)
        
        # Mida "una mica més gran" que 3.0 -> 3.3 polzades
        # I "sense que creixi l'espai de la capçalera" -> Això es controla amb els marges de la secció, no aquí.
        header_logo_width = Inches(3.3) 
        
        if target_logo:
             # Ajustar "amunt": El marge superior de la secció controla la posició base.
             # Aquí assegurem que no hi ha espai extra.
            run.add_picture(target_logo, width=header_logo_width)
        else:
            print("ERROR: Logo definitiu no trobat! No es mostrarà cap logo.")
            
    except Exception as e:
        print(f"ERROR LOGO HEADER: {e}")

def add_logo_to_body(doc):
    """Afegeix el logotip en gran al cos del document (per a la portada)"""
    try:
        p = doc.add_paragraph()
        p.alignment = Align.CENTER
        run = p.add_run()
        
        # PRIORITAT ABSOLUTA: logo_definitive.jpg (o png)
        logo_path_jpg = STATIC_DIR / "logo_definitive.jpg"
        logo_path_png = STATIC_DIR / "logo_definitive.png"
        
        target_logo = None
        if logo_path_jpg.exists(): target_logo = str(logo_path_jpg)
        elif logo_path_png.exists(): target_logo = str(logo_path_png)
        
        # Mida gran per portada
        cover_logo_width = Inches(4.5) 
        
        if target_logo:
            run.add_picture(target_logo, width=cover_logo_width)
            # Versió eliminada de la portada per petició: "no hace falta poner la version"
        else:
            print("ERROR: Logo definitiu no trobat! No es mostrarà cap logo.")
            
    except Exception as e:
        print(f"ERROR LOGO BODY: {e}")


def resize_to_box(img, max_w, max_h, allow_upscale=False):
    w, h = img.size
    ratio = min(max_w / w, max_h / h)
    if not allow_upscale and ratio > 1: return img
    new_w, new_h = int(w * ratio), int(h * ratio)
    return img.resize((new_w, new_h), Image.LANCZOS)

def to_jpeg_path(path):
    return path.with_suffix(".jpg")

def add_photo_block(doc, img_buffer, photo_num, page_w_cm, max_h_cm, description=""):
    # page_w_cm: Ample TOTAL disponible a la pàgina (per al text)
    # max_h_cm: Alçada màxima de la FOTO
    
    # 1. Preparar imatge (mida més gran per estètica, ex: 16.5cm)
    photo_display_w = 16.5 
    
    img_buffer.seek(0)
    with Image.open(img_buffer) as im:
        w, h = im.size; aspect_ratio = w / h if h else 1
        # Calculem dimensions finals
        final_w = photo_display_w
        final_h = final_w / aspect_ratio
        if final_h > max_h_cm:
            final_h = max_h_cm
            final_w = final_h * aspect_ratio
            
    # 2. Crear Taula contenidora (100% ample pàgina)
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = Align.CENTER
    # Forçar ample taula al total disponible
    tbl.autofit = False
    tbl.allow_autofit = False
    
    # Ajustar ample cel·la manualment
    cell = tbl.cell(0, 0)
    cell.width = Cm(page_w_cm)
    
    # --- EVITAR TALLS DE PÀGINA ---
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:cantSplit')
    trPr.append(trHeight)
    
    # 3. Inserir Imatge (Centrada)
    p_img = cell.paragraphs[0]; p_img.alignment = Align.CENTER
    p_img.paragraph_format.keep_with_next = True
    p_img.paragraph_format.space_before = Pt(0); p_img.paragraph_format.space_after = Pt(4)
    
    img_buffer.seek(0)
    run_img = p_img.add_run()
    run_img.add_picture(img_buffer, width=Cm(final_w), height=Cm(final_h))
    
    # 4. Inserir Text (Justificat i ample complet)
    if description:
        p_caption = cell.add_paragraph()
        p_caption.alignment = Align.JUSTIFY
        # Keep lines together
        p_caption.paragraph_format.keep_together = True
        
        # Mida de font dinàmica per encabir-ho tot
        # Si el text és llarg, reduïm la lletra perquè no salti de pàgina
        desc_len = len(description)
        font_size = 9
        if desc_len > 600: font_size = 8
        if desc_len > 900: font_size = 7
        
        # Format compacte
        r1 = p_caption.add_run(f"Fotografia núm. {photo_num}: ")
        r1.bold = True; r1.font.size = Pt(font_size); r1.font.name = 'Arial'
        r2 = p_caption.add_run(description)
        r2.font.size = Pt(font_size); r2.font.name = 'Arial'
    else:
        p_caption = cell.add_paragraph()
        p_caption.alignment = Align.CENTER
        p_caption.paragraph_format.keep_together = True
        r1 = p_caption.add_run(f"Fotografia núm. {photo_num}")
        r1.bold = True; r1.font.size = Pt(9); r1.font.name = 'Arial'

# --- GENERATION ---
def init_google_ai():
    api_key = os.environ.get("GEMINI_API_KEY", "AIzaSyCLz6Rvi-6MGV9tsdS3XTkMeTOQefJi6qw")
    genai.configure(api_key=api_key)
    print("DEBUG: Google AI (AI Studio) inicialitzat amb clau API.")

def init_vertex_ai():
    if IS_PROD:
        try:
            vertexai.init(project=os.environ.get('PROJECT_ID', 'infofoto-vector-art'), location="us-central1")
        except Exception as e:
            print(f"ERROR: Vertex AI init failed: {e}")

import traceback

def generate_ai_descriptions(evolucio: str, image_files: List[str], sid: str = None) -> Dict[str, str]:
    if not image_files: return {}
    
    print(f"DEBUG: GEN - Iniciant generació Studio per {len(image_files)} imatges. Relat: {bool(evolucio)}")
    init_google_ai()
    
    # Si no hi ha evolució, usem un prompt per defecte per no cancel·lar la generació
    context_evolucio = evolucio if evolucio else "Investigació d'accident de trànsit (sense relat previ)."
    
    # Models update: Restaurada prioritat a Gemini 3 Flash (Preview) per petició d'usuari
    # Fallback a 2.0 i 1.5 per seguretat
    models_to_try = ["gemini-3-flash-preview", "gemini-2.0-flash", "gemini-1.5-flash"]
    
    prompt_text = (
        "Ets un Cap d'Unitat d'Investigació i Reconstrucció d'Accidents de Trànsit (UIRAT) dels Mossos d'Esquadra. "
        "La teva feina és realitzar l'anàlisi tècnica d'un informe fotogràfic forense per al Jutjat.\n\n"
        f"CONTEXT DE L'ACCIDENT: '{context_evolucio}'.\n\n"
        "NORMATIVA DE REDACCIÓ:\n"
        "1. TONALITAT: Sobri, imparcial, eminentment tècnic i professional.\n"
        "2. VOCABULARI ESPECIALITZAT: Obligatori utilitzar conceptes de física i mecànica d'accidents (ex: deformació plàstica residual, transferència energetica d'impacte, eixos de compressió, petjades de frenada/lliscament, pèrdua d'integritat de l'habitacle, danys per abrasió, punts de conflicte primaris i secundaris).\n"
        "3. INTEGRACIÓ: Utilitza les dades del context per identificar els vehicles correctament (marca, model, matrícula si es coneix) i situar l'acció en el lloc descrit.\n"
        "4. EXCLUSIÓ: Prohibit descriure personal o vehicles d'emergència (bombers, SEM). Centra't en les evidències materials de l'accident.\n"
        "4. EXCLUSIÓ: Prohibit descriure personal o vehicles d'emergència (bombers, SEM). Centra't en les evidències materials de l'accident.\n"
        "5. EXTENSIÓ I RELLEVÀNCIA: \"Alargar un poco mas el texto, un poco más centrandose en factores relevantes\". Redacta de 4 a 6 línies denses explicant la mecànica, els danys i la física de l'accident. Prioritza informació útil per al Jutge.\n\n"
        "FORMAT: Retorna EXCLUSIVAMENT un objecte JSON amb la clau del fitxer i el valor amb la descripció."
    )
    
    contents = [prompt_text]
    # ... (resta de codi igual per preparar imatges)
    valid_images_count = 0
    total_imgs = len(image_files)
    
    for idx, fname in enumerate(image_files):
        update_status(f"Processant imatges ({idx+1}/{total_imgs})...", sid)
        try:
            # 1. Comprovar localment
            src_path = None
            for folder in [EDITED_DIR, MASTER_DIR, WORK_DIR]:
                if (folder/fname).exists():
                    src_path = folder/fname
                    break
            
            # Baixada de seguretat de GCS
            if not src_path and IS_PROD:
                for folder_name in ["work", "master", "edited"]:
                    p_target = MASTER_DIR / fname
                    print(f"DEBUG: GEN - Intentant descarregar {fname} de {folder_name}...")
                    if storage_download(f"uploads/{folder_name}/{fname}", p_target):
                        src_path = p_target
                        print(f"DEBUG: GEN - Descarregada correctament de {folder_name}")
                        break
            
            if src_path:
                with Image.open(src_path) as img:
                    if img.mode not in ("RGB", "L"): img = img.convert("RGB")
                    # Reduir una mica qualitat per anar ràpid a l'API
                    img_small = resize_to_box(img, 1024, 768) 
                    buf = io.BytesIO()
                    img_small.save(buf, format="JPEG", quality=85)
                    buf.seek(0)
                    
                    contents.append(f"NOM DEL FITXER: {fname}")
                    contents.append({
                        "mime_type": "image/jpeg",
                        "data": buf.read()
                    })
                    valid_images_count += 1
        except Exception as e:
             print(f"DEBUG: GEN Error imatge {fname}: {e}")
        
    if valid_images_count == 0:
        return {f: "Error: No s'han pogut llegir les imatges." for f in image_files}

    # Bucle de reintents
    for model_name in models_to_try:
        try:
            # "No hace falta ponerlo ni el modelo que se utiliza tampoco": Canviat a "Analitzant..." pur i dur.
            update_status(f"Analitzant imatges i generant descripcions...", sid)
            print(f"DEBUG: GEN - Intentant amb {model_name}...")
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(
                contents, 
                generation_config={
                    "response_mime_type": "application/json",
                    "temperature": 0.1
                }
            )
            if response and response.text:
                print(f"DEBUG: GEN - Èxit amb {model_name}. Processant JSON...")
                try:
                    raw_data = json.loads(response.text)
                except:
                    print(f"DEBUG: GEN - JSON Invàlid: {response.text}")
                    continue

                final_descs = {}
                
                # Estratègia 1: Diccionari directe o llista de dicts
                temp_dict = {}
                if isinstance(raw_data, list):
                    for item in raw_data:
                        if isinstance(item, dict): temp_dict.update(item)
                elif isinstance(raw_data, dict):
                    temp_dict = raw_data
                
                # Estratègia 2: Matching
                # Intentem casar les claus retornades amb els fitxers reals
                ai_keys = list(temp_dict.keys())
                print(f"DEBUG: GEN Claus Retornades: {ai_keys}")
                print(f"DEBUG: Imatges Esperades: {image_files}")
                
                for i, fname in enumerate(image_files):
                    # 2a. Coincidència exacta
                    if fname in temp_dict:
                        final_descs[fname] = temp_dict[fname]
                    # 2b. Coincidència parcial (si la IA es deixa .jpg)
                    elif fname.split('.')[0] in temp_dict:
                         final_descs[fname] = temp_dict[fname.split('.')[0]]
                    # 2c. Fallback per índex (si la IA retorna claus rares però en ordre)
                    elif i < len(ai_keys):
                        # Assumim que l'ordre es respecta
                        key_at_index = ai_keys[i]
                        final_descs[fname] = temp_dict[key_at_index]
                
                # VALIDACIÓ DE LA RESPOSTA
                # Si hem aconseguit casar alguna cosa, ho donem per bo.
                # Si el diccionari està buit, considerem que el model ha fallat i provem el següent.
                if final_descs:
                    print(f"DEBUG: Èxit rotund amb {model_name}. {len(final_descs)} descripcions generades.")
                    return final_descs
                else:
                    print(f"ADVERTÈNCIA: Model {model_name} ha retornat JSON vàlid però cap coincidència amb fitxers. Provant següent model...")
                    continue # Saltem al següent model

            else:
                 print(f"ADVERTÈNCIA: Model {model_name} resposta buida. Provant següent...")
                 continue

        except Exception as e:
            msg = f"ADVERTÈNCIA: Model {model_name} ha fallat: {e}"
            print(msg)
            # traceback.print_exc() # Opcional per no embrutar logs
            continue 
            
    # Si arribem aquí, MALA SORT: Cap model ha funcionat.
    return {f: "Error crític: Tots els models de generació han fallat. Intenta-ho de nou." for f in image_files}

# --- ROUTES ---

@app.get("/")
def index():
    # Eliminat esborrat global de GCS: era un coll d'ampolla crític
    session.clear()
    return render_template("index.html")

@app.post("/start_session")
def start_session():
    print(f"DEBUG: start_session START - Path: {request.path}")
    try:
        session.clear()
        sid = str(uuid.uuid4())
        session['sid'] = sid
        print(f"DEBUG: start_session NEW SID: {sid}")
        
        data = {
            'nat': request.form.get("nat","").strip(),
            'dil': request.form.get("dil","").strip(),
            'tip1': request.form.get("tip1","").strip(),
            'tip2': request.form.get("tip2","").strip(),
            'jutjat': request.form.get("jutjat", "").strip(),
            'localitat': request.form.get("localitat", "").strip(),
            'qualitat': request.form.get("qualitat", "atenea").strip(),
            'latest_uploads': [],
            'image_order': [],
            'created_at': time.time()
        }
        
        # Guardar (Sincronitzat amb GCS per estabilitat)
        save_gcs_session(data, sid=sid)
        print(f"DEBUG: start_session SUCCESS - SID: {sid}")
        
        return jsonify(ok=True, sid=sid)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"ERROR: start_session failed: {e}")
        return jsonify(ok=False, error=str(e)), 500

@app.post("/finalize_upload")
def finalize_upload():
    # Forçar sessió si ve per URL
    sid_param = request.args.get('sid')
    if sid_param: session['sid'] = sid_param
    
    req_data = request.get_json()
    new_files = req_data.get('files', [])
    sdata = load_gcs_session()
    current_uploads = sdata.get('latest_uploads', [])
    current_order = sdata.get('image_order', [])
    
    updated_uploads = current_uploads + [f for f in new_files if f not in current_uploads]
    if not current_order:
        updated_order = updated_uploads
    else:
        updated_order = current_order + [f for f in new_files if f not in current_order]
        
    update_gcs_session({'latest_uploads': updated_uploads, 'image_order': updated_order})
    return jsonify(ok=True)

import threading

@app.post("/upload")
def upload():
    # Forçar sessió si ve per URL
    sid_param = request.args.get('sid')
    if sid_param: session['sid'] = sid_param

    files = request.files.getlist("photos")
    work_filename = None
    for f in files:
        if not (f and f.filename): continue
        clean_name = secure_filename(Path(f.filename).stem)
        unique_id = int(time.time() * 1000) % 1000000
        base_name = f"{clean_name}_{unique_id}"
        
        tmp_path = WORK_DIR / (base_name + "_tmp"); f.save(tmp_path)
        try:
            with Image.open(tmp_path) as img:
                if img.mode not in ("RGB", "L"): img = img.convert("RGB")
                
                # Editor quality & AI (Sincronitzat)
                work_img = resize_to_box(img, 1360, 768, allow_upscale=False)
                work_path = to_jpeg_path(WORK_DIR / base_name); work_img.save(work_path, format="JPEG", quality=90, optimize=True)
                if IS_PROD:
                    storage_save(work_path, f"uploads/work/{work_path.name}")
                
                # Report quality (Background)
                master_img = resize_to_box(img, 1920, 1080, allow_upscale=False)
                master_path = to_jpeg_path(MASTER_DIR / base_name); master_img.save(master_path, format="JPEG", quality=90, optimize=True)
                if IS_PROD:
                    run_in_bg(storage_save, master_path, f"uploads/master/{master_path.name}")
                
                work_filename = work_path.name
        except Exception as e:
            print(f"[UPLOAD] Error processant {f.filename}: {e}")
            return jsonify(ok=False, error=str(e)), 500
        finally:
            try: tmp_path.unlink(missing_ok=True)
            except Exception: pass
            
    return jsonify(ok=True, filename=work_filename)

@app.get("/order")
def order():
    try:
        # Forçar sessió si ve per URL
        sid_param = request.args.get('sid')
        if sid_param: session['sid'] = sid_param

        sdata = load_gcs_session(sid_param)
        images_to_display = sdata.get('image_order') or sdata.get('latest_uploads', [])
        
        return render_template("order.html", 
                               images=images_to_display, 
                               nat=sdata.get('nat',''), 
                               dil=sdata.get('dil',''), 
                               tip1=sdata.get('tip1',''), 
                               tip2=sdata.get('tip2',''), 
                               jutjat=sdata.get('jutjat',''), 
                               localitat=sdata.get('localitat',''), 
                               evolucio=sdata.get('evolucio', ''),
                               image_descriptions=sdata.get('image_descriptions', {}),
                               ts=int(time.time()))
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"S'ha produït un error carregant la llista d'imatges: {str(e)}", 500

@app.post("/update_order")
def update_order():
    data = request.get_json()
    update_gcs_session({'image_order': data.get('order', [])})
    return jsonify(ok=True)

@app.post("/delete/<path:filename>")
def delete_image(filename):
    for folder in (EDITED_DIR, WORK_DIR, MASTER_DIR):
        (folder / Path(filename).name).unlink(missing_ok=True)
    
    sdata = load_gcs_session()
    current_order = sdata.get('image_order', [])
    updates = {}
    if current_order:
        updates['image_order'] = [img for img in current_order if img != filename]
        
    current_uploads = sdata.get('latest_uploads', [])
    if current_uploads:
         updates['latest_uploads'] = [img for img in current_uploads if img != filename]
    
    if updates: update_gcs_session(updates)
    return jsonify(ok=True)

@app.get("/uploads/<path:filename>")
def uploaded_file(filename):
    p_edit = EDITED_DIR / filename
    p_work = WORK_DIR / filename
    
    if IS_PROD:
        if not p_edit.exists(): storage_download(f"uploads/edited/{filename}", p_edit)
        if not p_work.exists() and not p_edit.exists(): storage_download(f"uploads/work/{filename}", p_work)

    directory_to_check = EDITED_DIR if p_edit.exists() else WORK_DIR
    response = make_response(send_from_directory(directory_to_check, filename))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'; response.headers['Expires'] = '0'
    return response

@app.get("/edit/<path:filename>")
def edit(filename): return render_template("edit.html", filename=filename, ts=int(time.time()))

@app.post("/save_edit/<path:filename>")
def save_edit(filename):
    f = request.files.get("file")
    if not f: return jsonify(ok=False, error="No s'ha rebut cap arxiu.")
    EDITED_DIR.mkdir(parents=True, exist_ok=True)
    out_path = EDITED_DIR / Path(secure_filename(filename)).name
    f.save(out_path)
    storage_save(out_path, f"uploads/edited/{out_path.name}")
    return jsonify(ok=True, path=str(out_path))

@app.post("/generate_api")
def generate_api():
    form_sid = request.form.get('sid')
    if form_sid: session['sid'] = form_sid
    
    sdata = load_gcs_session()
    evolucio = request.form.get("evolucio", "").strip()
    order_csv = request.form.get("order","").strip()
    images_ordered = [x for x in (y.strip() for y in order_csv.split(",")) if x]
    if not images_ordered: images_ordered = sdata.get('image_order', [])
    
    # Recollir descripcions manuals si n'hi ha
    manual_descriptions = {}
    for key in request.form:
        if key.startswith("desc_"): manual_descriptions[key[5:]] = request.form[key]
    
    # Guardar estat inicial
    sdata.update({'evolucio': evolucio, 'image_order': images_ordered, 'status_message': "Iniciant tasca de fons..."})
    save_gcs_session(sdata)

    def bg_task(sid, evol, imgs, manuals):
        try:
            print(f"DEBUG: BG_TASK Iniciada per {sid}")
            ai_descs = generate_ai_descriptions(evol, imgs, sid=sid)
            
            # Recarregar sdata per evitar sobreescriure si ha canviat
            current_sdata = load_gcs_session(sid)
            final_descs = manuals.copy()
            for fname, desc in ai_descs.items():
                if not final_descs.get(fname):
                    final_descs[fname] = desc
            
            current_sdata['image_descriptions'] = final_descs
            current_sdata['status_message'] = "[READY]"
            save_gcs_session(current_sdata, sid)
            print(f"DEBUG: BG_TASK Finalitzada per {sid}")
        except Exception as e:
            print(f"ERROR BG_TASK: {e}")
            cs = load_gcs_session(sid)
            cs['status_message'] = f"[ERROR] {str(e)}"
            save_gcs_session(cs, sid)

    run_in_bg(bg_task, session.get('sid'), evolucio, images_ordered, manual_descriptions)
    return jsonify(ok=True)

@app.get("/review")
def review_report():
    # Forçar sessió si ve per URL
    sid_param = request.args.get('sid')
    if sid_param: session['sid'] = sid_param

    sdata = load_gcs_session()
    images_ordered = sdata.get('image_order', [])
    descriptions = sdata.get('image_descriptions', {})
    
    if not images_ordered:
        return redirect(url_for('order'))
        
    return render_template("review.html", 
                           images=images_ordered, 
                           descriptions=descriptions,
                           sid=session.get('sid', ''))

@app.post("/save_descriptions")
def save_descriptions():
    form_sid = request.form.get('sid')
    if form_sid: session['sid'] = form_sid
    sdata = load_gcs_session()
    
    final_descriptions = {}
    for key, value in request.form.items():
        if key.startswith("desc_"):
            fname = key[5:]
            final_descriptions[fname] = value.strip()
            
    sdata['image_descriptions'] = final_descriptions
    save_gcs_session(sdata)
    return jsonify(ok=True)

@app.post("/create_report")
def create_report():
    try:
        # INTENTAR RECUPERAR SESSIÓ VIA FORMULARI (MÉS FIABLE QUE COOKIE)
        form_sid = request.form.get('sid')
        if form_sid:
            print(f"DEBUG: Forçant sessió amb SID del formulari: {form_sid}")
            session['sid'] = form_sid
            
        sdata = load_gcs_session()
        
        images_ordered = sdata.get('image_order', [])
        if not images_ordered:
            return "Error: No s'han trobat imatges per generar l'informe.", 400

        nat_code = sdata.get('nat', 'SENSE DADES')
        dil_code = sdata.get('dil', 'SENSE DADES')
        tip1 = sdata.get('tip1', '')
        tip2 = sdata.get('tip2', '')
        jutjat = sdata.get('jutjat', 'Jutjat desconegut')
        localitat = sdata.get('localitat', 'Localitat desconeguda')
        qualitat = sdata.get('qualitat', 'atenea')
        
        # Recollir descripcions
        final_descriptions = {}
        for key, value in request.form.items():
            if key.startswith("desc_"):
                fname = key[5:]
                final_descriptions[fname] = value.strip()
            
        # Actualitzar sessió (optimitzat: evitem reload)
        sdata['image_descriptions'] = final_descriptions
        save_gcs_session(sdata)

        # Configuració de qualitat
        target_w, target_h, allow_up, jpg_quality = (2560, 2560, False, 95) if qualitat == 'vector' else (1920, 1080, True, 80)
        
        processed_images = []; valid_images_ordered = []
        
        print(f"DEBUG: Iniciant processament de {len(images_ordered)} imatges.")
        
        for idx, name in enumerate(images_ordered):
            print(f"DEBUG: Processant imatge {idx+1}/{len(images_ordered)}: {name}")
            try:
                p_edit, p_master, p_work = EDITED_DIR/name, MASTER_DIR/name, WORK_DIR/name
                
                # Descarregar si és necessari (amb retry implícit per lògica visual)
                if IS_PROD:
                    # Optimització: Només baixar el que necessitem si ja existeix
                    found = False
                    if p_edit.exists(): found = True
                    elif storage_download(f"uploads/edited/{name}", p_edit): found = True
                    
                    if not found:
                         if p_master.exists(): found = True
                         elif storage_download(f"uploads/master/{name}", p_master): found = True
                    
                    if not found:
                         if not p_work.exists(): storage_download(f"uploads/work/{name}", p_work)
    
                src_path = p_edit if p_edit.exists() else (p_master if p_master.exists() else (p_work if p_work.exists() else None))
                
                if src_path:
                    with Image.open(src_path) as img:
                        if img.mode in ('RGBA', 'P'): img = img.convert('RGB')
                        img2 = resize_to_box(img, target_w, target_h, allow_upscale=allow_up)
                        buffer = io.BytesIO()
                        img2.save(buffer, format='JPEG', quality=jpg_quality, optimize=True)
                        buffer.seek(0)
                        processed_images.append(buffer)
                        valid_images_ordered.append(name)
                else:
                    print(f"ADVERTÈNCIA: No s'ha trobat cap arxiu font per a {name}")
            except Exception as e:
                print(f"ERROR processant imatge {name}: {e}")
                
        images_ordered = valid_images_ordered
        if not images_ordered:
             return "Error: Totes les imatges han fallat al processar-se.", 500
        
        print("DEBUG: Generant document DOCX...")
        doc = Document()
        # ... (rest of document generation logic stays similar but ensured inside try) ...
        # [REPEATING DOX GENERATION LOGIC TO ENSURE SCOPE IS CORRECT]
        # Marges "Estirats" (més estrets per aprofitar espai)
        margin_cm = 1.0 
        page_h_cm, page_w_cm = 29.7, 21.0
        
        sec_titol = doc.sections[0]
        sec_titol.different_first_page_header_footer = False
        sec_titol.page_height = Cm(page_h_cm); sec_titol.page_width = Cm(page_w_cm)
        sec_titol.top_margin = Cm(margin_cm); sec_titol.bottom_margin = Cm(margin_cm)
        sec_titol.left_margin = Cm(margin_cm); sec_titol.right_margin = Cm(margin_cm)
        
        # 1. Logo Capçalera (Petit) - NOMÉS definim aquí la capçalera principal (pàgines següents)
        # La portada NO tindrà logo a la capçalera (perquè ja en té un de gran)
        # Per fer-ho, NO cridem add_logo_to_header aquí, sinó que confiem en el loop principal per a les seccions noves.
        # Però la primera secció (sec_titol) és la portada. 
        # Simplement NO fem add_logo_to_header(sec_titol.header)
        
        create_footer(sec_titol, nat_code, dil_code, qualitat)
        
        # 2. Logo Portada (Gran)
        add_logo_to_body(doc)
        
        doc.add_paragraph("\n"); p = doc.add_paragraph("Informe Fotogràfic"); p.alignment = Align.CENTER
        p.runs[0].font.size = Pt(33); p.runs[0].font.name = 'Arial'
        doc.add_paragraph() # Espai doble
        
        unit_name = "Unitat d'Investigació d'Accidents de Trànsit" if qualitat == 'atenea' else "Unitat d'Atestats de Trànsit"
        
        lines = []
        if nat_code or dil_code:
            parts = []
            if qualitat == 'vector':
                if dil_code: parts.append(f"Dil. {dil_code}")
                if nat_code: parts.append(f"NAT {nat_code}")
            else:
                if nat_code: parts.append(f"NAT {nat_code}")
                if dil_code: parts.append(f"Dil. {dil_code}")
            lines.append(" - ".join(parts))
            
        lines.extend(["Àrea Regional de Trànsit Metropolitana Nord", unit_name, f"Data d’emissió de l’informe: {datetime.date.today().strftime('%d/%m/%Y')}"])
        if tip1 or tip2:
            label = "Núm. de TIP:" if (tip1 and not tip2) or (tip2 and not tip1) else "Núms. de TIP:"
            lines.append(f"{label} {' - '.join([x for x in [tip1, tip2] if x])}")
        for txt in lines:
            p = doc.add_paragraph(txt); p.alignment = Align.CENTER
            p.runs[0].font.size = Pt(21); p.runs[0].font.name = 'Arial'
            doc.add_paragraph() # Espai doble

        # Destinació al final de la portada (Protocol oficial)
        doc.add_paragraph("\n")
        p_dest = doc.add_paragraph(f"S'adreça al {jutjat} de la localitat de {localitat}"); p_dest.alignment = Align.CENTER
        # "La ultima linea la del juzgado un poco más perqueña (2 puntos)" -> 21 - 2 = 19
        p_dest.runs[0].font.size = Pt(19); p_dest.runs[0].font.name = 'Arial'
        doc.add_paragraph()
        
        # Càlcul d'espai útil
        usable_w_cm = page_w_cm - (margin_cm * 2)
        # Alçada pàgina (29.7) - Marges (2) - Header/Footer (3.5) = ~24.2cm
        # Per 2 fotos+text en una pàgina: 9.0cm és el mínim exigit per l'usuari.
        photo_max_h = 9.0 
        
        photo_counter = 1
        
        # Iterar de 2 en 2
        import math
        total_photos = len(processed_images)
        chunks = [processed_images[i:i + 2] for i in range(0, len(processed_images), 2)]
        images_names_chunks = [images_ordered[i:i + 2] for i in range(0, len(images_ordered), 2)]
        
        for i, chunk in enumerate(chunks):
            # Nova pàgina per a cada parell de fotos
            sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
            sec.different_first_page_header_footer = False
            sec.page_height=Cm(page_h_cm); sec.page_width=Cm(page_w_cm)
            sec.top_margin=Cm(margin_cm); sec.bottom_margin=Cm(margin_cm)
            sec.left_margin=Cm(margin_cm); sec.right_margin=Cm(margin_cm)
            
            # Header i Footer a cada pagina
            add_logo_to_header(sec.header)
            create_footer(sec, nat_code, dil_code, qualitat)
            set_vertical_alignment(sec, 'top') # Alineació superior
            
            # Processar les 1 o 2 fotos d'aquest chunk
            current_names = images_names_chunks[i]
            
            for j, img_data in enumerate(chunk):
                f_name = current_names[j]
                desc = final_descriptions.get(f_name, "")
                
                # Afegir bloc de foto
                add_photo_block(doc, img_data, photo_counter, usable_w_cm, photo_max_h, desc)
                photo_counter += 1
                
                # Espaiador entre fotos (només si n'hi ha una altra després en la mateixa pàgina)
                if j == 0 and len(chunk) > 1:
                    p_sep = doc.add_paragraph()
                    # Reduïm espai entre fotos per guanyar marge
                    p_sep.paragraph_format.space_before = Pt(2)
                    p_sep.paragraph_format.space_after = Pt(2)
                
        # Diligència final
        sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
        sec.different_first_page_header_footer = False
        sec.page_height=Cm(page_h_cm); sec.page_width=Cm(page_w_cm)
        add_logo_to_header(sec.header); sec.footer.is_linked_to_previous=True
        sec.top_margin=Cm(margin_cm); sec.bottom_margin=Cm(margin_cm)
        sec.left_margin=Cm(margin_cm); sec.right_margin=Cm(margin_cm)
        # Alineació vertical superior per a la diligència final
        set_vertical_alignment(sec, 'top')
            
        doc.add_paragraph(); doc.add_paragraph()
        hora = datetime.datetime.now().strftime('%H:%M'); dia = datetime.date.today().strftime('%d/%m/%Y')
        p = doc.add_paragraph(); p.alignment = Align.JUSTIFY
        titol = "Diligència de Tramesa d'Informe Fotogràfic:"
        r = p.add_run(titol); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(12)
        
        # Calcular total de pàgines reals (Portada + Pàgines de fotos + Pàgina de tancament)
        photo_pages = (len(processed_images) + 1) // 2
        total_pages = 1 + photo_pages + 1
        
        p_txt = f" Que a les {hora} del dia {dia}, es finalitza aquest Informe fotogràfic, el qual consta de {len(images_ordered)} fotografies i un total de {total_pages} pàgines. S'adreça al {jutjat} de la localitat de {localitat}."
        r_txt = p.add_run(p_txt); r_txt.font.name = 'Arial'; r_txt.font.size = Pt(12)
        p_c = doc.add_paragraph(); r_c = p_c.add_run("Perque Consti ho Certifico"); r_c.bold=True; r_c.font.size=Pt(12); r_c.font.name = 'Arial'
    
        # Generació Directa en Memòria (Motor Pro)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_nat = nat_code.replace('/', '_') if nat_code else "SENSE_NAT"
        rname_download = f'Informe_{safe_nat}.docx'
        
        # Backup a GCS si som a PROD
        if IS_PROD:
            try:
                bucket = get_bucket()
                if bucket:
                    unique_id = int(time.time())
                    blob = bucket.blob(f"reports/Informe_{safe_nat}_{unique_id}.docx")
                    blob.upload_from_string(file_stream.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    print(f"BACKUP_OK: {safe_nat}_{unique_id}")
            except Exception as e:
                print(f"ADVERTÈNCIA: Error backup GCS: {e}")
    
        # Retornar el fitxer amb les millors capçaleres de streaming possibles
        response = make_response(send_file(
            file_stream,
            as_attachment=True,
            download_name=rname_download,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ))
        
        # Forçar capçaleres per evitar bloquejos de descàrrega
        response.headers["Content-Disposition"] = f"attachment; filename=\"{rname_download}\""
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"ERROR_CRITIC_DOWNLOAD: {e}")
        return f"Error en la generació del document: {str(e)}", 500

@app.get("/status")
def get_status():
    sdata = load_gcs_session()
    return jsonify(status=sdata.get('status_message', 'Processant...'))

def update_status(message, sid=None):
    print(f"STATUS UPDATE: {message}")
    # Si no passem sid, cridarà get_sid() que fallarà sense context
    # És responsabilitat del caller passar el sid si està en background
    update_gcs_session({'status_message': message}, sid=sid)



if __name__ == "__main__":
    default_port = 8080 if IS_PROD else 5051
    port = int(os.environ.get("PORT", default_port))
    app.run(host="0.0.0.0", port=port, debug=not IS_PROD)
