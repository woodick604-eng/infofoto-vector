# app.py — Dual workflow: 1360×768 (work) + 1920×1080 (report)
import io, time, datetime, os
from pathlib import Path
import sys
from typing import List
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, jsonify, make_response, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as Align
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, UnidentifiedImageError

# Firebase imports
import firebase_admin
from firebase_admin import credentials, firestore, storage

def resource_path(relative_path):
    try:
        base_path = Path(sys._MEIPASS)  # type: ignore[attr-defined]
    except Exception:
        base_path = Path(__file__).resolve().parent
    return base_path / relative_path

# Configuració Firebase

IS_PROD = os.environ.get("K_SERVICE") is not None
if not firebase_admin._apps:
    if IS_PROD:
        firebase_admin.initialize_app(options={'storageBucket': f"{os.environ.get('PROJECT_ID', 'infofoto-vector-art')}.appspot.com"})
    else:
        # Localment no inicialitzem si no tenim credencials, o usem emuladors
        try:
            firebase_admin.initialize_app()
        except Exception:
            pass

db = firestore.client(database_id='infofotovector') if firebase_admin._apps else None

def get_bucket():

    if not firebase_admin._apps: return None
    try: return storage.bucket()
    except Exception: return None

BASE_DIR    = resource_path(".")
UPLOAD_DIR  = BASE_DIR / "uploads"
MASTER_DIR  = UPLOAD_DIR / "master"
WORK_DIR    = UPLOAD_DIR / "work"
EDITED_DIR  = UPLOAD_DIR / "edited"
REPORTS_DIR = BASE_DIR / "reports"
STATIC_DIR  = BASE_DIR / "static"
LOGO_PATH   = STATIC_DIR / "mossos.jpg"

for d in (UPLOAD_DIR, MASTER_DIR, WORK_DIR, EDITED_DIR, REPORTS_DIR, STATIC_DIR):
    d.mkdir(parents=True, exist_ok=True)

# Helper functions for Storage
def storage_save(local_path: Path, storage_rel_path: str):
    if not IS_PROD: return
    bucket = get_bucket()
    if not bucket: return
    blob = bucket.blob(storage_rel_path)
    blob.upload_from_filename(str(local_path))

def storage_download(storage_rel_path: str, local_path: Path):
    if not IS_PROD: return False
    bucket = get_bucket()
    if not bucket: return False
    blob = bucket.blob(storage_rel_path)
    if blob.exists():
        blob.download_to_filename(str(local_path))
        return True
    return False

app = Flask(__name__, template_folder=str(BASE_DIR/"templates"), static_folder=str(BASE_DIR/"static"))
app.secret_key = os.environ.get("SECRET_KEY", "una-clau-secreta-molt-dificil-de-endevinar")
app.config['GOOGLE_CLIENT_ID'] = os.environ.get("GOOGLE_CLIENT_ID", "814718439112-2hcqqhsbbb2b67btpcqgtepakhmkhkkk.apps.googleusercontent.com")


def resize_to_box(img: Image.Image, max_w: int, max_h: int, allow_upscale: bool=False) -> Image.Image:
    w, h = img.size
    ratio = min(max_w / w, max_h / h)
    if not allow_upscale and ratio >= 1.0:
        return img
    new_size = (max(1, int(w * ratio)), max(1, int(h * ratio)))
    return img.resize(new_size, Image.Resampling.LANCZOS)

def to_jpeg_path(path: Path) -> Path:
    return path.with_suffix(".jpg")

def add_field(p, c):
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
    i  = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = c; run._r.append(i)
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); run._r.append(f2)

def create_footer(sec, nat_code: str, dil_code: str = "", qualitat: str = "atenea"):
    ft = sec.footer
    ft.is_linked_to_previous = False
    if ft.paragraphs: ft.paragraphs[0]._element.getparent().remove(ft.paragraphs[0]._element)
    parts = []
    if qualitat == 'vector':
        if dil_code: parts.append(f"Dil. {dil_code}")
        if nat_code: parts.append(f"NAT {nat_code}")
    else:
        if nat_code: parts.append(f"NAT {nat_code}")
        if dil_code: parts.append(f"Dil. {dil_code}")
    nat_dil_string = " - ".join(parts)
    final_parts = []
    if nat_dil_string: final_parts.append(nat_dil_string)
    final_parts.append("ART MNORD")
    left_text = " ".join(final_parts)
    tbl = ft.add_table(1, 2, width=Inches(6))
    t0, t1 = tbl.rows[0].cells
    r = t0.paragraphs[0].add_run(left_text); r.bold = True; r.font.size = Pt(10); r.font.name = "Arial"
    p1 = t1.paragraphs[0]; p1.alignment = Align.RIGHT
    p1.add_run("Pàgina ").font.name="Arial"; add_field(p1, 'PAGE'); p1.add_run(" de ").font.name="Arial"; add_field(p1, 'NUMPAGES')

def add_logo_to_header(header):
    try:
        p = header.paragraphs[0]; p.clear(); p.alignment = Align.LEFT
        run = p.add_run(); run.add_picture(str(LOGO_PATH), width=Inches(3.2))
    except Exception as e:
        print(f"ADVERTÈNCIA: No s'ha pogut carregar el logo: {e}")

def add_photo_block(doc, img_buffer, photo_num, max_w_cm, max_h_cm):
    img_buffer.seek(0)
    with Image.open(img_buffer) as im:
        w, h = im.size; aspect_ratio = w / h if h else 1
        img_w_cm = max_w_cm; img_h_cm = img_w_cm / aspect_ratio
        if img_h_cm > max_h_cm: img_h_cm = max_h_cm; img_w_cm = img_h_cm * aspect_ratio
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = Align.CENTER
    img_cell = tbl.cell(0, 0)
    tc = img_cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(int(img_w_cm * 567))); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
    p_img = img_cell.paragraphs[0]; p_img.alignment = Align.CENTER
    p_img.paragraph_format.space_before = Pt(0); p_img.paragraph_format.space_after = Pt(0)
    img_buffer.seek(0); p_img.add_run().add_picture(img_buffer, width=Cm(img_w_cm), height=Cm(img_h_cm))
    p_caption = doc.add_paragraph(); p_caption.alignment = Align.CENTER
    r1 = p_caption.add_run(f"Fotografia núm. {photo_num}"); r1.bold = True; r1.font.size = Pt(14); r1.font.name = 'Arial'
    p_caption.add_run(": ").font.size = Pt(12); p_caption.add_run("").font.size = Pt(12)

@app.get("/")
def index():
    session.clear()
    for folder in [EDITED_DIR]:
        if not folder.exists(): continue
        for item in folder.iterdir():
            if item.is_file() and not item.name.startswith('.'):
                try: item.unlink()
                except Exception as e: print(f"No s'ha pogut esborrar {item}: {e}")
    return render_template("index.html")

@app.post("/upload")
def upload():
    session['nat'] = request.form.get("nat","").strip()
    session['dil'] = request.form.get("dil","").strip()
    session['tip1'] = request.form.get("tip1","").strip()
    session['tip2'] = request.form.get("tip2","").strip()
    session['jutjat'] = request.form.get("jutjat", "").strip()
    session['localitat'] = request.form.get("localitat", "").strip()
    session['qualitat'] = request.form.get("qualitat", "atenea").strip()
    files = request.files.getlist("photos"); uploaded_names = []
    for f in files:
        if not (f and f.filename): continue
        base_name = secure_filename(Path(f.filename).stem)
        tmp_path = WORK_DIR / (base_name + "_tmp"); f.save(tmp_path)
        try:
            with Image.open(tmp_path) as img:
                if img.mode not in ("RGB", "L"): img = img.convert("RGB")
                master_img = resize_to_box(img, 1920, 1080, allow_upscale=False)
                master_path = to_jpeg_path(MASTER_DIR / base_name); master_img.save(master_path, format="JPEG", quality=90, optimize=True)
                storage_save(master_path, f"uploads/master/{master_path.name}")
                
                work_img = resize_to_box(img, 1360, 768, allow_upscale=False)
                work_path = to_jpeg_path(WORK_DIR / base_name); work_img.save(work_path, format="JPEG", quality=90, optimize=True)
                storage_save(work_path, f"uploads/work/{work_path.name}")
                
                uploaded_names.append(work_path.name)
        except Exception as e:
            print(f"No s'ha pogut processar {f.filename}: {e}")
        finally:
            try: tmp_path.unlink(missing_ok=True)
            except Exception: pass
    session['latest_uploads'] = uploaded_names; session.pop('image_order', None)
    return redirect(url_for("order"))

@app.get("/order")
def order():
    nat = session.get('nat', ''); dil = session.get('dil', ''); tip1 = session.get('tip1', ''); tip2 = session.get('tip2', '')
    jutjat = session.get('jutjat', ''); localitat = session.get('localitat', '')
    current_files = {p.name for p in WORK_DIR.iterdir() if p.is_file() and not p.name.startswith('.')}
    if 'image_order' in session:
        images_to_display = [img for img in session['image_order'] if img in current_files or (EDITED_DIR / img).exists()]
    else:
        images_to_display = sorted(session.get('latest_uploads', [])); session['image_order'] = images_to_display
    return render_template("order.html", images=images_to_display, nat=nat, dil=dil, tip1=tip1, tip2=tip2, jutjat=jutjat, localitat=localitat, ts=int(time.time()))

@app.post("/update_order")
def update_order():
    data = request.get_json(); session['image_order'] = data.get('order', []); return jsonify(ok=True)

@app.post("/delete/<path:filename>")
def delete_image(filename):
    for folder in (EDITED_DIR, WORK_DIR, MASTER_DIR):
        (folder / Path(filename).name).unlink(missing_ok=True)
    session['image_order'] = [img for img in session.get('image_order', []) if img != filename]
    return jsonify(ok=True)

@app.get("/uploads/<path:filename>")
def uploaded_file(filename):
    p_edit = EDITED_DIR / filename
    p_work = WORK_DIR / filename
    
    if IS_PROD:
        if not p_edit.exists(): storage_download(f"uploads/edited/{filename}", p_edit)
        if not p_work.exists() and not p_edit.exists(): storage_download(f"uploads/work/{filename}", p_work)

    directory_to_check = EDITED_DIR if p_edit.exists() else WORK_DIR
    response = make_response(send_from_directory(directory_to_check, filename))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'; response.headers['Expires'] = '0'
    return response

@app.get("/edit/<path:filename>")
def edit(filename): return render_template("edit.html", filename=filename, ts=int(time.time()))

@app.post("/save_edit/<path:filename>")
def save_edit(filename):
    f = request.files.get("file")
    if not f: return jsonify(ok=False, error="No s'ha rebut cap arxiu.")
    EDITED_DIR.mkdir(parents=True, exist_ok=True)
    out_path = EDITED_DIR / Path(secure_filename(filename)).name
    f.save(out_path)
    storage_save(out_path, f"uploads/edited/{out_path.name}")
    return jsonify(ok=True, path=str(out_path))

@app.post("/generate")
def generate():
    nat_code = session.get('nat', ''); dil_code = session.get('dil', '')
    tip1 = session.get('tip1', ''); tip2 = session.get('tip2', '')
    jutjat = session.get('jutjat', ''); localitat = session.get('localitat', '')
    qualitat = session.get('qualitat', 'atenea')
    order_csv = request.form.get("order","").strip()
    images_ordered = [x for x in (y.strip() for y in order_csv.split(",")) if x]
    if not images_ordered: return redirect(url_for('order'))
    if qualitat == 'vector':
        target_w, target_h, allow_up, jpg_quality = 2560, 2560, False, 95
    else:
        target_w, target_h, allow_up, jpg_quality = 1920, 1080, True, 85
    processed_images: List[io.BytesIO] = []; valid_images_ordered: List[str] = []
    for name in images_ordered:
        try:
            p_edit = EDITED_DIR / name
            p_master = MASTER_DIR / name
            p_work = WORK_DIR / name
            
            if IS_PROD:
                if not p_edit.exists(): storage_download(f"uploads/edited/{name}", p_edit)
                if not p_edit.exists() and not p_master.exists(): storage_download(f"uploads/master/{name}", p_master)
                if not p_edit.exists() and not p_master.exists() and not p_work.exists(): storage_download(f"uploads/work/{name}", p_work)

            if p_edit.exists(): src_path = p_edit
            else:
                src_path = p_master
                if not src_path.exists():
                    src_path = p_work if p_work.exists() else src_path
            
            with Image.open(src_path) as img:
                if img.mode in ('RGBA', 'P'): img = img.convert('RGB')
                img2 = resize_to_box(img, target_w, target_h, allow_upscale=allow_up)
                buffer = io.BytesIO(); img2.save(buffer, format='JPEG', quality=jpg_quality, optimize=True); buffer.seek(0)
                processed_images.append(buffer); valid_images_ordered.append(name)
        except (FileNotFoundError, UnidentifiedImageError) as e:
            print(f"ADVERTÈNCIA: L'arxiu '{name}' no s'ha trobat o està malmès. S'ignorarà. Error: {e}")
    images_ordered = valid_images_ordered
    if not images_ordered: return "Error: Cap imatge vàlida per processar.", 400
    photo_pages_count = 0; i = 0
    while i < len(images_ordered):
        img1_buffer = processed_images[i]; img1_buffer.seek(0)
        with Image.open(img1_buffer) as img1: w1, h1 = img1.size
        is_vertical1 = h1 > w1; is_next_horizontal = False
        if i + 1 < len(images_ordered):
            img2_buffer = processed_images[i+1]; img2_buffer.seek(0)
            with Image.open(img2_buffer) as img2: w2, h2 = img2.size
            if h2 <= w2: is_next_horizontal = True
        if not is_vertical1 and is_next_horizontal: i += 2
        else: i += 1
        photo_pages_count += 1
    total_pages = 1 + photo_pages_count
    if jutjat and localitat: total_pages += 1
    doc = Document()
    page_h_cm_val, page_w_cm_val, margin_cm_val = 29.7, 21.0, 2.5
    sec_titol = doc.sections[0]
    sec_titol.page_height = Cm(page_h_cm_val); sec_titol.page_width = Cm(page_w_cm_val)
    sec_titol.top_margin = Cm(margin_cm_val); sec_titol.bottom_margin = Cm(margin_cm_val)
    sec_titol.left_margin = Cm(margin_cm_val); sec_titol.right_margin = Cm(margin_cm_val)
    add_logo_to_header(sec_titol.header); create_footer(sec_titol, nat_code, dil_code, qualitat)
    doc.add_paragraph("\n\n\n"); p = doc.add_paragraph("Informe Fotogràfic"); p.alignment = Align.CENTER
    p.runs[0].font.size = Pt(33); p.runs[0].font.name = 'Arial'
    lines = []
    if nat_code or dil_code:
        line = []
        if qualitat == 'vector':
            if dil_code: line.append(f"Dil. {dil_code}")
            if nat_code: line.append(f"NAT {nat_code}")
        else:
            if nat_code: line.append(f"NAT {nat_code}")
            if dil_code: line.append(f"Dil. {dil_code}")
        lines.append(" - ".join(line))
    lines.extend(["Àrea Regional de Trànsit Metropolitana Nord","Unitat d'Investigació d'Accidents de Trànsit",f"Data d’emissió de l’informe: {datetime.date.today().strftime('%d/%m/%Y')}"])
    if tip1 or tip2:
        label = "Núm. de TIP:" if (tip1 and not tip2) or (tip2 and not tip1) else "Núms. de TIP:"
        parts = [x for x in [tip1, tip2] if x]
        lines.append(f"{label} {' - '.join(parts)}")
    for txt in lines:
        doc.add_paragraph(); p = doc.add_paragraph(txt); p.alignment = Align.CENTER
        p.runs[0].font.size = Pt(21); p.runs[0].font.name = 'Arial'
    usable_w_cm = page_w_cm_val - (margin_cm_val * 2); usable_h_cm = page_h_cm_val - (margin_cm_val * 2)
    photo_counter = 1; i = 0
    while i < len(processed_images):
        sec_foto = doc.add_section(WD_SECTION_START.NEW_PAGE)
        sec_foto.page_height=sec_titol.page_height; sec_foto.page_width=sec_titol.page_width
        sec_foto.top_margin=sec_titol.top_margin; sec_foto.bottom_margin=sec_titol.bottom_margin
        sec_foto.left_margin=sec_titol.left_margin; sec_foto.right_margin=sec_titol.right_margin
        sectPr=sec_foto._sectPr; vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); sectPr.append(vAlign)
        sec_foto.header.is_linked_to_previous=True; create_footer(sec_foto, nat_code, dil_code, qualitat)
        img1_buffer = processed_images[i]; img1_buffer.seek(0)
        with Image.open(img1_buffer) as img1: w1, h1 = img1.size
        is_vertical1 = h1 > w1; is_next_horizontal = False
        if i + 1 < len(processed_images):
            img2_buffer = processed_images[i+1]; img2_buffer.seek(0)
            with Image.open(img2_buffer) as img2: w2, h2 = img2.size
            if h2 <= w2: is_next_horizontal = True
        if not is_vertical1 and is_next_horizontal:
            add_photo_block(doc, processed_images[i], photo_counter, usable_w_cm, (usable_h_cm - 2.5*2 - 1)/2); photo_counter += 1
            doc.add_paragraph()
            add_photo_block(doc, processed_images[i+1], photo_counter, usable_w_cm, (usable_h_cm - 2.5*2 - 1)/2); photo_counter += 1
            i += 2
        else:
            add_photo_block(doc, processed_images[i], photo_counter, usable_w_cm, usable_h_cm - 3.0); photo_counter += 1
            i += 1
    if jutjat and localitat:
        sec_final = doc.add_section(WD_SECTION_START.NEW_PAGE)
        sec_final.page_height=doc.sections[-2].page_height; sec_final.page_width=doc.sections[-2].page_width
        sec_final.left_margin=doc.sections[-2].left_margin; sec_final.right_margin=doc.sections[-2].right_margin
        sec_final.top_margin=doc.sections[-2].top_margin; sec_final.bottom_margin=doc.sections[-2].bottom_margin
        sec_final.header.is_linked_to_previous=True; sec_final.footer.is_linked_to_previous=True
        sectPr=sec_final._sectPr; vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'top'); sectPr.append(vAlign)
        doc.add_paragraph(); doc.add_paragraph()
        hora_actual = datetime.datetime.now().strftime('%H:%M'); dia_actual = datetime.date.today().strftime('%d/%m/%Y')
        num_fotos = len(images_ordered)
        p_dil = doc.add_paragraph(); p_dil.alignment = Align.JUSTIFY
        if qualitat == 'vector':
            run_titol = p_dil.add_run("Diligència de Tramesa d'Informe Fotogràfic:"); run_titol.bold = True; run_titol.font.name = 'Arial'; run_titol.font.size = Pt(14)
            run_text = p_dil.add_run(f" Que a les {hora_actual} del dia {dia_actual}, es finalitza aquest Informe fotogràfic amb número de diligències {dil_code}, que consta de {num_fotos} fotografies i {total_pages} pàgines, i es tramet al {jutjat} de la localitat de {localitat}."); run_text.font.name = 'Arial'; run_text.font.size = Pt(14)
        else:
            run_titol = p_dil.add_run("Diligència d'Informe Fotogràfic:"); run_titol.bold = True; run_titol.font.name = 'Arial'; run_titol.font.size = Pt(14)
            run_text = p_dil.add_run(f" Que a les {hora_actual} del dia {dia_actual}, es finalitza aquest Informe fotogràfic que consta de {num_fotos} fotografies i {total_pages} pàgines, i es tramet al {jutjat} de la localitat de {localitat}."); run_text.font.name = 'Arial'; run_text.font.size = Pt(14)
        p_cert = doc.add_paragraph(); cert_run = p_cert.add_run("Consti i Certifico"); cert_run.bold = True; cert_run.font.name = 'Arial'; cert_run.font.size = Pt(14)
    safe_nat = nat_code.replace('/', '_') if nat_code else "SENSE_NAT"
    report_name = f'Informe_{safe_nat}_{int(time.time())}.docx'
    report_path = REPORTS_DIR / report_name; doc.save(report_path)
    storage_save(report_path, f"reports/{report_name}")
    return redirect(url_for('download_report', filename=report_name))

@app.route('/report/<path:filename>')
def download_report(filename):
    if IS_PROD:
        bucket = get_bucket()
        if not bucket: return "Storage error", 500
        blob = bucket.blob(f"reports/{filename}")
        if not blob.exists():
            return "Report not found", 404
        
        file_stream = io.BytesIO()
        blob.download_to_file(file_stream)
        file_stream.seek(0)
        
        return send_from_directory(REPORTS_DIR, filename, as_attachment=True) if (REPORTS_DIR/filename).exists() else make_response(file_stream.read(), 200, {'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'Content-Disposition': f'attachment; filename="{filename}"'})
    else:
        return send_from_directory(REPORTS_DIR, filename, as_attachment=True)

if __name__ == "__main__":
    default_port = 8080 if IS_PROD else 5051
    port = int(os.environ.get("PORT", default_port))
    app.run(host="0.0.0.0", port=port, debug=not IS_PROD)